"""Legal hallucination checking routes."""
import base64
import os
import time
import uuid
from io import BytesIO

import httpx
from fastapi import APIRouter, HTTPException, Request

from db import get_user_by_auth_token


router = APIRouter(prefix="/api/hallucination", tags=["hallucination"])

YUANDIAN_HALL_DETECT_URL = "https://open.chineselaw.com/open/hall_detect"
MAX_TEXT_LENGTH = 50000
MAX_UPLOAD_BYTES = 30 * 1024 * 1024


async def require_user(request: Request) -> dict:
    user = await get_user_by_auth_token(request.cookies.get("user_token"))
    if not user:
        raise HTTPException(status_code=401, detail="请先登录")
    return user


def _extract_docx_text(file_data: str) -> str:
    from docx import Document

    raw = base64.b64decode(file_data)
    doc = Document(BytesIO(raw))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_pdf_text(file_data: str, max_pages: int = 80) -> str:
    from pypdf import PdfReader

    raw = base64.b64decode(file_data)
    reader = PdfReader(BytesIO(raw))
    parts = []

    for page_index, page in enumerate(reader.pages[:max_pages], 1):
        text = (page.extract_text() or "").strip()
        if text:
            parts.append(f"--- 第 {page_index} 页 ---\n{text}")

    if len(reader.pages) > max_pages:
        parts.append(f"（PDF 共 {len(reader.pages)} 页，仅提取前 {max_pages} 页）")

    return "\n\n".join(parts)


def _extract_txt_text(file_data: str) -> str:
    raw = base64.b64decode(file_data)
    for encoding in ("utf-8", "gb18030", "gbk"):
        try:
            return raw.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace").strip()


@router.post("/extract-file")
async def extract_hallucination_file(request: Request):
    """Extract plain text from uploaded Word/PDF/TXT for hall detection."""
    await require_user(request)

    try:
        body = await request.json()
        file_info = body.get("file") or {}
        filename = str(file_info.get("name") or "").strip()
        mime_type = str(file_info.get("mime_type") or "").strip().lower()
        file_data = str(file_info.get("data") or "")
    except Exception:
        raise HTTPException(status_code=400, detail="无效的请求")

    if not file_data:
        raise HTTPException(status_code=400, detail="请上传 Word、PDF 或 TXT 文件")

    try:
        raw_size = len(base64.b64decode(file_data, validate=True))
    except Exception:
        raise HTTPException(status_code=400, detail="文件内容格式不正确")

    if raw_size > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=400, detail="文件超过 30MB，请压缩后再上传")

    lower_name = filename.lower()
    try:
        if lower_name.endswith(".docx") or "wordprocessingml" in mime_type:
            text = _extract_docx_text(file_data)
        elif lower_name.endswith(".pdf") or mime_type == "application/pdf":
            text = _extract_pdf_text(file_data)
        elif lower_name.endswith(".txt") or mime_type.startswith("text/"):
            text = _extract_txt_text(file_data)
        else:
            raise HTTPException(status_code=400, detail="仅支持 DOCX、PDF、TXT 文件")
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"文件解析失败: {exc}")

    if not text:
        raise HTTPException(status_code=400, detail="未能从文件中提取到文本")

    return {
        "success": True,
        "filename": filename,
        "text": text[:MAX_TEXT_LENGTH],
        "truncated": len(text) > MAX_TEXT_LENGTH,
        "char_count": min(len(text), MAX_TEXT_LENGTH),
    }


@router.post("/check")
async def check_hallucination(request: Request):
    """Proxy Yuandian hall_detect without exposing the API key to the browser."""
    await require_user(request)

    try:
        body = await request.json()
        text = str(body.get("text") or "").strip()
    except Exception:
        raise HTTPException(status_code=400, detail="无效的请求")

    if not text:
        raise HTTPException(status_code=400, detail="请输入需要校验的文本")
    if len(text) > MAX_TEXT_LENGTH:
        raise HTTPException(status_code=400, detail="文本过长，请控制在 5 万字以内分段校验")

    api_key = os.getenv("YUANDIAN_API_KEY", "").strip()
    if not api_key:
        raise HTTPException(status_code=503, detail="元典 API Key 未配置，无法进行幻觉校验")

    request_id = f"hall_{int(time.time() * 1000)}_{uuid.uuid4().hex[:8]}"
    headers = {
        "Content-Type": "application/json; charset=utf-8",
        "Accept": "application/json",
        "X-API-Key": api_key,
        "X-Request-ID": request_id,
    }

    try:
        async with httpx.AsyncClient(timeout=httpx.Timeout(120.0, connect=15.0)) as client:
            response = await client.post(
                YUANDIAN_HALL_DETECT_URL,
                headers=headers,
                json={"text": text},
            )
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="元典幻觉校验接口超时，请稍后重试")
    except httpx.RequestError as exc:
        raise HTTPException(status_code=502, detail=f"元典幻觉校验接口请求失败: {exc}")

    response_request_id = response.headers.get("X-Request-ID") or request_id
    try:
        payload = response.json()
    except Exception:
        payload = {"raw": response.text}

    if response.status_code >= 400:
        message = "元典幻觉校验失败"
        if isinstance(payload, dict):
            message = payload.get("message") or payload.get("detail") or payload.get("error") or message
        raise HTTPException(
            status_code=response.status_code if response.status_code < 600 else 502,
            detail=f"{message}（request_id: {response_request_id}）",
        )

    return {
        "success": True,
        "request_id": response_request_id,
        "result": payload,
    }
