"""
AI 法律助手 — 后端主入口
FastAPI 应用、路由、中间件
"""
import json
import os
import asyncio
import base64
import posixpath
import re
import time
import uuid
import zipfile
from io import BytesIO
from pathlib import Path

from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles

from prompts import get_prompt
from usage import extract_usage, record_response_usage, record_usage

# 加载环境变量
env_path = Path(__file__).parent / ".env"
load_dotenv(env_path)

# 配置
ACCESS_PASSWORD = os.getenv("ACCESS_PASSWORD")
if not ACCESS_PASSWORD:
    raise RuntimeError("ACCESS_PASSWORD 未配置，请在 backend/.env 中设置访问密码")

SESSION_MAX_AGE_HOURS = int(os.getenv("SESSION_MAX_AGE_HOURS", "24"))
SESSION_MAX_AGE = SESSION_MAX_AGE_HOURS * 3600

# 数据库配置
DB_HOST = os.getenv("DB_HOST")
if not DB_HOST:
    raise RuntimeError("DB_HOST 未配置，请在 backend/.env 中设置数据库地址")
DB_PORT = int(os.getenv("DB_PORT", "3306"))
DB_USER = os.getenv("DB_USER")
DB_PASSWORD = os.getenv("DB_PASSWORD")
DB_NAME = os.getenv("DB_NAME", "ai_law_helper")
if not DB_USER or DB_PASSWORD is None:
    raise RuntimeError("DB_USER/DB_PASSWORD 未配置，请在 backend/.env 中设置数据库账号")

# 创建 FastAPI 应用
app = FastAPI(
    title="AI 法律助手",
    description="一个温馨的法律AI助手网站",
    version="1.0.0",
)


# ===== 启动事件 =====
@app.on_event("startup")
async def startup_event():
    """初始化数据库 + 加载 MCP 工具"""
    from db import init_db, bootstrap_admin_user
    await init_db(host=DB_HOST, user=DB_USER, password=DB_PASSWORD, database=DB_NAME, port=DB_PORT)
    from usage import ensure_usage_tables
    await ensure_usage_tables()
    # 创建默认管理员，并把历史会话迁移到该用户下
    await bootstrap_admin_user(os.getenv("ADMIN_USERNAME", "loveHmt"), ACCESS_PASSWORD)

    from mcp_client import init_mcp, refresh_tools
    YUANDIAN_API_KEY = os.getenv("YUANDIAN_API_KEY", "")
    ZHIPU_API_KEY = os.getenv("ZHIPU_API_KEY", "")
    YUANDIAN_MCP_ENABLED = os.getenv("YUANDIAN_MCP_ENABLED", "true").lower() == "true"
    init_mcp(api_key=YUANDIAN_API_KEY, zhipu_api_key=ZHIPU_API_KEY, enabled=YUANDIAN_MCP_ENABLED)
    await refresh_tools()


@app.on_event("shutdown")
async def shutdown_event():
    """关闭数据库连接 + MCP 持久化连接"""
    from db import close_db
    await close_db()
    from mcp_client import close_mcp_connections
    await close_mcp_connections()


# CORS 配置
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

from routes.hallucination import router as hallucination_router
from routes.usage import router as usage_router
from routes.translate import router as translate_router
from routes.admin import router as admin_router

app.include_router(hallucination_router)
app.include_router(usage_router)
app.include_router(translate_router)
app.include_router(admin_router)

# ===== 静态文件服务 =====
frontend_path = Path(__file__).parent.parent / "frontend"

for subdir in ["css", "js", "assets"]:
    sub_path = frontend_path / subdir
    if sub_path.exists():
        app.mount(f"/{subdir}", StaticFiles(directory=str(sub_path)), name=f"static-{subdir}")


@app.get("/")
async def serve_index():
    return FileResponse(frontend_path / "index.html")


@app.get("/admin")
async def serve_admin():
    return FileResponse(frontend_path / "admin.html")


@app.get("/mobile")
async def serve_mobile():
    return FileResponse(frontend_path / "mobile.html")


@app.get("/document")
async def serve_document_workbench():
    return FileResponse(frontend_path / "document.html")


@app.get("/research")
async def serve_research_workbench():
    return FileResponse(frontend_path / "research.html")


@app.get("/translate")
async def serve_translate_page():
    return FileResponse(frontend_path / "translate.html")


@app.get("/usage")
async def serve_usage_workbench():
    return FileResponse(frontend_path / "usage.html")


def extract_docx_text(file_data: str) -> str:
    """从前端传来的 base64 docx 中提取段落和表格文本。"""
    from docx import Document

    raw = base64.b64decode(file_data)
    doc = Document(BytesIO(raw))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_pdf_text(file_data: str, max_pages: int = 80) -> str:
    """从前端传来的 base64 PDF 中提取文本型 PDF 内容。"""
    from pypdf import PdfReader

    raw = base64.b64decode(file_data)
    reader = PdfReader(BytesIO(raw))
    parts = []

    for page_index, page in enumerate(reader.pages[:max_pages], 1):
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            parts.append(f"--- 第 {page_index} 页 ---\n{text}")

    if len(reader.pages) > max_pages:
        parts.append(f"（PDF 共 {len(reader.pages)} 页，仅提取前 {max_pages} 页）")

    return "\n\n".join(parts)


def _read_mineru_zip(zip_bytes: bytes) -> dict:
    """读取 MinerU 结果包，返回 Markdown、图片资源和结构 JSON。"""
    image_exts = (".png", ".jpg", ".jpeg", ".webp")
    image_mimes = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
    }
    with zipfile.ZipFile(BytesIO(zip_bytes)) as zf:
        names = zf.namelist()
        markdown_names = [
            name for name in names
            if name.lower().endswith((".md", ".markdown"))
        ]
        if not markdown_names:
            raise RuntimeError("MinerU 结果包中未找到 Markdown 文件")
        markdown_names.sort(key=lambda name: ("/" in name, len(name)))

        assets = {}
        basename_assets = {}
        for name in names:
            lower = name.lower()
            ext = next((e for e in image_exts if lower.endswith(e)), "")
            if not ext:
                continue
            data = zf.read(name)
            if len(data) > 256 * 1024 or len(assets) >= 12:
                continue
            uri = f"data:{image_mimes[ext]};base64,{base64.b64encode(data).decode('ascii')}"
            normalized = posixpath.normpath(name)
            assets[normalized] = uri
            basename_assets[posixpath.basename(normalized)] = uri

        json_files = {}
        for name in names:
            lower = name.lower()
            if not lower.endswith(".json"):
                continue
            try:
                json_files[posixpath.basename(name)] = json.loads(zf.read(name).decode("utf-8", errors="replace"))
            except Exception:
                continue

        parts = []
        for name in markdown_names:
            text = zf.read(name).decode("utf-8", errors="replace").strip()
            if not text:
                continue
            md_dir = posixpath.dirname(posixpath.normpath(name))

            def replace_image(match: re.Match) -> str:
                alt = match.group(1)
                src = (match.group(2) or "").strip()
                if src.startswith(("http://", "https://", "data:")):
                    return match.group(0)
                normalized_src = posixpath.normpath(posixpath.join(md_dir, src))
                data_uri = assets.get(normalized_src) or assets.get(posixpath.normpath(src)) or basename_assets.get(posixpath.basename(src))
                if not data_uri:
                    return match.group(0)
                return f"![{alt}]({data_uri})"

            text = re.sub(r"!\[([^\]]*)\]\(([^)]+)\)", replace_image, text)
            parts.append(text)

    return {
        "source": "mineru",
        "markdown": "\n\n".join(parts),
        "asset_count": len(assets),
        "json_files": json_files,
    }


def extract_document_with_mineru_package(file_data: str, filename: str, verify_ssl: bool = True) -> dict:
    """通过 MinerU API 解析文档，返回结果包中的 Markdown 和可展示资源。"""
    import httpx

    token = os.getenv("MINERU_API_TOKEN", "").strip()
    if not token:
        raise RuntimeError("MINERU_API_TOKEN 未配置")

    base_url = os.getenv("MINERU_API_BASE_URL", "https://mineru.net").rstrip("/")
    model_version = os.getenv("MINERU_MODEL_VERSION", "vlm")
    poll_interval = float(os.getenv("MINERU_POLL_INTERVAL_SECONDS", "3"))
    poll_timeout = int(os.getenv("MINERU_POLL_TIMEOUT_SECONDS", "180"))

    raw = base64.b64decode(file_data)
    safe_name = filename or f"{uuid.uuid4().hex}.pdf"
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json",
    }

    with httpx.Client(timeout=60, verify=verify_ssl) as client:
        upload_resp = client.post(
            f"{base_url}/api/v4/file-urls/batch",
            headers=headers,
            json={
                "enable_formula": True,
                "enable_table": True,
                "language": "ch",
                "model_version": model_version,
                "files": [
                    {
                        "name": safe_name,
                        "is_ocr": True,
                        "data_id": uuid.uuid4().hex,
                    }
                ],
            },
        )
        upload_resp.raise_for_status()
        upload_data = upload_resp.json()
        if upload_data.get("code") not in (0, 200, "0"):
            raise RuntimeError(upload_data.get("msg") or f"MinerU 创建任务失败: {upload_data}")

        data = upload_data.get("data") or {}
        batch_id = data.get("batch_id")
        file_urls = data.get("file_urls") or []
        upload_url = file_urls[0] if file_urls else ""
        if not batch_id or not upload_url:
            raise RuntimeError(f"MinerU 未返回上传地址: {upload_data}")

        put_resp = client.put(upload_url, content=raw)
        put_resp.raise_for_status()

        deadline = time.time() + poll_timeout
        result_data = None
        while time.time() < deadline:
            result_resp = client.get(
                f"{base_url}/api/v4/extract-results/batch/{batch_id}",
                headers={"Authorization": f"Bearer {token}"},
            )
            result_resp.raise_for_status()
            result_json = result_resp.json()
            if result_json.get("code") not in (0, 200, "0"):
                raise RuntimeError(result_json.get("msg") or f"MinerU 查询任务失败: {result_json}")
            result_data = result_json.get("data") or {}
            extract_results = result_data.get("extract_result") or []
            if extract_results:
                state = extract_results[0].get("state", "").lower()
                if state in ("done", "success", "completed"):
                    break
                if state in ("failed", "error"):
                    raise RuntimeError(extract_results[0].get("err_msg") or "MinerU 解析失败")
            time.sleep(poll_interval)
        else:
            raise TimeoutError("MinerU 解析超时")

        extract_results = (result_data or {}).get("extract_result") or []
        full_zip_url = extract_results[0].get("full_zip_url") if extract_results else ""
        if not full_zip_url:
            raise RuntimeError(f"MinerU 未返回结果下载地址: {result_data}")

        zip_resp = client.get(full_zip_url)
        zip_resp.raise_for_status()

    return _read_mineru_zip(zip_resp.content)


def extract_pdf_with_mineru_package(file_data: str, filename: str, verify_ssl: bool = True) -> dict:
    """通过 MinerU API 解析 PDF，返回结果包中的 Markdown 和可展示资源。"""
    return extract_document_with_mineru_package(file_data, filename, verify_ssl=verify_ssl)


def extract_docx_with_mineru_package(file_data: str, filename: str, verify_ssl: bool = True) -> dict:
    """通过 MinerU API 解析 Word，返回结果包中的 Markdown 和可展示资源。"""
    return extract_document_with_mineru_package(file_data, filename, verify_ssl=verify_ssl)


def extract_pdf_text_with_mineru(file_data: str, filename: str, verify_ssl: bool = True) -> str:
    """通过 MinerU API 解析 PDF，返回结果 zip 中的 Markdown 文本。"""
    return extract_pdf_with_mineru_package(file_data, filename, verify_ssl=verify_ssl)["markdown"]


def extract_docx_text_with_mineru(file_data: str, filename: str, verify_ssl: bool = True) -> str:
    """通过 MinerU API 解析 Word，返回结果 zip 中的 Markdown 文本。"""
    return extract_docx_with_mineru_package(file_data, filename, verify_ssl=verify_ssl)["markdown"]


def extract_pdf_text_best_effort(file_data: str, filename: str) -> str:
    """优先使用 MinerU，失败时降级到本地文本型 PDF 提取。"""
    mineru_error = ""
    if os.getenv("MINERU_API_TOKEN", "").strip():
        try:
            text = extract_pdf_text_with_mineru(file_data, filename)
            if text.strip():
                return text
        except Exception as e:
            mineru_error = str(e)
            retry_insecure = os.getenv("MINERU_RETRY_INSECURE_SSL", "true").lower() == "true"
            ssl_error_markers = ("ssl", "certificate", "cert", "tls", "handshake")
            if retry_insecure and any(marker in mineru_error.lower() for marker in ssl_error_markers):
                try:
                    text = extract_pdf_text_with_mineru(file_data, filename, verify_ssl=False)
                    if text.strip():
                        return text
                except Exception as retry_error:
                    mineru_error = f"{mineru_error}; SSL 兼容重试失败：{retry_error}"

    text = extract_pdf_text(file_data)
    if text.strip():
        if mineru_error:
            return f"（MinerU 解析失败，已使用本地 PDF 文本提取降级：{mineru_error}）\n\n{text}"
        return text

    if mineru_error:
        raise RuntimeError(f"MinerU 解析失败，且本地未提取到文本：{mineru_error}")
    return ""


def extract_docx_text_best_effort(file_data: str, filename: str) -> str:
    """优先使用 MinerU 解析 Word，失败时降级到 python-docx。"""
    mineru_error = ""
    if os.getenv("MINERU_API_TOKEN", "").strip():
        try:
            text = extract_docx_text_with_mineru(file_data, filename)
            if text.strip():
                return text
        except Exception as e:
            mineru_error = str(e)
            retry_insecure = os.getenv("MINERU_RETRY_INSECURE_SSL", "true").lower() == "true"
            ssl_error_markers = ("ssl", "certificate", "cert", "tls", "handshake")
            if retry_insecure and any(marker in mineru_error.lower() for marker in ssl_error_markers):
                try:
                    text = extract_docx_text_with_mineru(file_data, filename, verify_ssl=False)
                    if text.strip():
                        return text
                except Exception as retry_error:
                    mineru_error = f"{mineru_error}; SSL 兼容重试失败：{retry_error}"

    text = extract_docx_text(file_data)
    if text.strip():
        if mineru_error:
            return f"（MinerU 解析失败，已使用本地 Word 文本提取降级：{mineru_error}）\n\n{text}"
        return text

    if mineru_error:
        raise RuntimeError(f"MinerU 解析失败，且本地未提取到文本：{mineru_error}")
    return ""


def build_file_description(file_info: dict, max_chars: int = 20000) -> str:
    fname = file_info.get("name", "unknown")
    fmime = file_info.get("mime_type", "application/octet-stream")
    fdata = file_info.get("data", "")
    lower_name = fname.lower()

    if fmime.startswith("text/") or lower_name.endswith((".txt", ".md", ".csv", ".json", ".xml")):
        try:
            text_content = base64.b64decode(fdata).decode("utf-8", errors="replace")
            return f"[附件: {fname}]\n```\n{text_content[:max_chars]}\n```"
        except Exception:
            return f"[附件: {fname} ({fmime})，文件大小约 {len(fdata) * 3 // 4} 字节，文本读取失败]"

    if lower_name.endswith(".docx") or fmime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            text_content = extract_docx_text_best_effort(fdata, fname)
            if not text_content.strip():
                return f"[Word附件: {fname}]\n（未提取到可读文本，可能是扫描件或图片型文档）"
            return f"[Word附件: {fname}]\n```\n{text_content[:max_chars]}\n```"
        except Exception as e:
            return f"[Word附件: {fname}]\n（读取失败：{str(e)}）"

    if lower_name.endswith(".doc"):
        return f"[Word附件: {fname}]\n（暂不支持旧版 .doc 格式，请另存为 .docx 后上传）"

    if lower_name.endswith(".pdf") or fmime == "application/pdf":
        try:
            text_content = extract_pdf_text_best_effort(fdata, fname)
            if not text_content.strip():
                return f"[PDF附件: {fname}]\n（未提取到可读文本，可能是扫描件或图片型 PDF）"
            return f"[PDF附件: {fname}]\n```\n{text_content[:max_chars]}\n```"
        except Exception as e:
            return f"[PDF附件: {fname}]\n（读取失败：{str(e)}）"

    if fmime.startswith("image/"):
        return f"[图片附件: {fname}]\n![](data:{fmime};base64,{fdata[:5000]})"

    return f"[附件: {fname} ({fmime})]"


def summarize_uploaded_files(files: list[dict]) -> list[dict]:
    summaries = []
    for file_info in (files or [])[:8]:
        if not isinstance(file_info, dict):
            continue
        data = file_info.get("data") or ""
        size = file_info.get("size")
        if size is None and isinstance(data, str):
            size = len(data) * 3 // 4
        summaries.append({
            "name": file_info.get("name") or "unknown",
            "mime_type": file_info.get("mime_type") or file_info.get("type") or "",
            "size": size or 0,
        })
    return summaries


# ===== API 路由 =====

async def get_current_user(request: Request) -> dict | None:
    from db import get_user_by_auth_token
    return await get_user_by_auth_token(request.cookies.get("user_token"))


async def require_current_user(request: Request) -> dict:
    user = await get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="请先登录")
    return user


async def require_admin_user(request: Request) -> dict:
    user = await require_current_user(request)
    if not user.get("is_admin"):
        raise HTTPException(status_code=403, detail="需要管理员权限")
    return user


def public_user(user: dict) -> dict:
    return {
        "id": user["id"],
        "username": user["username"],
        "is_admin": bool(user["is_admin"]),
        "is_active": bool(user["is_active"]),
    }


@app.post("/api/login")
async def login(request: Request):
    """登录接口 — 返回当前活跃会话列表"""
    from db import verify_user, create_auth_session, create_session, list_sessions
    try:
        body = await request.json()
        username = body.get("username", os.getenv("ADMIN_USERNAME", "loveHmt")).strip()
        password = body.get("password", "")
    except Exception:
        username = os.getenv("ADMIN_USERNAME", "loveHmt")
        password = ""

    user = await verify_user(username, password)
    if not user:
        raise HTTPException(status_code=401, detail="用户名或密码错误")

    # 不自动创建新会话 — 如果已有会话就复用最后一个，否则才新建
    existing = await list_sessions(user["id"])
    if existing:
        token = existing[0]["id"]  # 最近的会话
    else:
        token = await create_session(user["id"])
        existing = await list_sessions(user["id"])
    user_token = await create_auth_session(user["id"], SESSION_MAX_AGE)

    response = JSONResponse({
        "success": True,
        "message": "欢迎回来！",
        "sessions": existing,  # 返回会话列表
        "current_session": token,
        "user": public_user(user),
    })
    response.set_cookie(
        key="user_token",
        value=user_token,
        max_age=SESSION_MAX_AGE,
        httponly=True,
        samesite="lax",
        secure=False,
        path="/",
    )
    response.set_cookie(
        key="session_token",
        value=token,
        max_age=SESSION_MAX_AGE,
        httponly=True,
        samesite="lax",
        secure=False,
        path="/",
    )
    return response


@app.get("/api/sessions")
async def get_sessions(request: Request):
    """获取所有会话列表"""
    from db import get_session, list_sessions, create_session
    user = await require_current_user(request)
    token = request.cookies.get("session_token")
    session = await get_session(token, user["id"]) if token else None
    sessions = await list_sessions(user["id"])
    if not session:
        if sessions:
            token = sessions[0]["id"]
        else:
            token = await create_session(user["id"])
            sessions = await list_sessions(user["id"])
    response = JSONResponse({"sessions": sessions, "current_session": token, "user": public_user(user)})
    response.set_cookie(
        key="session_token",
        value=token,
        max_age=SESSION_MAX_AGE,
        httponly=True,
        samesite="lax",
        secure=False,
        path="/",
    )
    return response


@app.post("/api/sessions/new")
async def new_session(request: Request):
    """创建新会话"""
    from db import create_session
    user = await require_current_user(request)

    try:
        body = await request.json()
        title = body.get("title", "新对话")
    except Exception:
        title = "新对话"

    new_token = await create_session(user["id"], title)
    response = JSONResponse({"success": True, "session_id": new_token})
    response.set_cookie(
        key="session_token",
        value=new_token,
        max_age=SESSION_MAX_AGE,
        httponly=True,
        samesite="lax",
        secure=False,
        path="/",
    )
    return response


@app.post("/api/sessions/switch")
async def switch_session(request: Request):
    """切换到指定会话"""
    from db import get_session
    user = await require_current_user(request)
    try:
        body = await request.json()
        target_id = body.get("session_id", "")
    except Exception:
        raise HTTPException(status_code=400, detail="无效的请求")

    if not target_id:
        raise HTTPException(status_code=400, detail="缺少 session_id")

    session = await get_session(target_id, user["id"])
    if not session:
        raise HTTPException(status_code=404, detail="会话不存在")

    response = JSONResponse({"success": True, "session_id": target_id})
    response.set_cookie(
        key="session_token",
        value=target_id,
        max_age=SESSION_MAX_AGE,
        httponly=True,
        samesite="lax",
        secure=False,
        path="/",
    )
    return response


@app.post("/api/sessions/rename")
async def rename_session(request: Request):
    """重命名会话"""
    from db import get_session, update_session_title
    user = await require_current_user(request)
    token = request.cookies.get("session_token")

    try:
        body = await request.json()
        session_id = body.get("session_id", token or "")
        title = body.get("title", "新对话")
    except Exception:
        raise HTTPException(status_code=400, detail="无效的请求")

    session = await get_session(session_id, user["id"]) if session_id else None
    if not session:
        raise HTTPException(status_code=401, detail="请先登录")

    await update_session_title(session_id, title, user["id"])
    return {"success": True}


@app.post("/api/sessions/delete")
async def delete_session_api(request: Request):
    """删除指定会话"""
    from db import get_session, delete_session, list_sessions, create_session
    user = await require_current_user(request)
    token = request.cookies.get("session_token")

    try:
        body = await request.json()
        session_id = body.get("session_id", "")
    except Exception:
        raise HTTPException(status_code=400, detail="无效的请求")

    if not session_id:
        raise HTTPException(status_code=400, detail="缺少 session_id")

    session = await get_session(session_id, user["id"])
    if not session:
        raise HTTPException(status_code=404, detail="会话不存在")

    await delete_session(session_id, user["id"])

    # 如果删除的是当前会话，自动切换到最新会话
    if session_id == token:
        existing = await list_sessions(user["id"])
        if existing:
            new_token = existing[0]["id"]
        else:
            new_token = await create_session(user["id"])
        response = JSONResponse({"success": True, "new_session": new_token})
        response.set_cookie(
            key="session_token", value=new_token,
            max_age=SESSION_MAX_AGE, httponly=True, samesite="lax", secure=False, path="/",
        )
        return response

    return {"success": True}


@app.post("/api/logout")
async def logout(request: Request):
    from db import delete_auth_session
    await delete_auth_session(request.cookies.get("user_token"))
    response = JSONResponse({"success": True})
    response.delete_cookie("user_token")
    response.delete_cookie("session_token")
    return response


@app.get("/api/check-auth")
async def check_auth(request: Request):
    user = await get_current_user(request)
    if not user:
        return {"authenticated": False}
    from db import get_session, list_sessions, create_session
    token = request.cookies.get("session_token")
    session = await get_session(token, user["id"]) if token else None
    sessions = await list_sessions(user["id"])
    if not session:
        if sessions:
            token = sessions[0]["id"]
        else:
            token = await create_session(user["id"])
    return {"authenticated": True, "user": public_user(user), "current_session": token}


@app.get("/api/messages")
async def get_current_messages(request: Request):
    from db import get_session, get_messages
    user = await require_current_user(request)
    token = request.cookies.get("session_token")
    session = await get_session(token, user["id"]) if token else None
    if not session:
        raise HTTPException(status_code=401, detail="请先登录")

    messages = await get_messages(token)
    visible_messages = [
        {"role": m["role"], "content": m.get("content") or ""}
        for m in messages
        if m.get("role") in ("user", "assistant") and m.get("content")
    ]
    return {"messages": visible_messages}


@app.post("/api/chat")
async def chat(request: Request):
    """对话接口 — SSE 流式输出，支持 Function Calling + MCP"""
    from db import (
        get_session,
        get_messages,
        save_new_messages,
        clear_messages,
        touch_session,
        update_session_title,
    )
    user = await require_current_user(request)
    token = request.cookies.get("session_token")
    session = await get_session(token, user["id"]) if token else None
    if not session:
        raise HTTPException(status_code=401, detail="请先登录")

    await touch_session(token, user["id"])

    try:
        body = await request.json()
        message = body.get("message", "")
        display_message = body.get("display_message") or message
        scene = body.get("scene", "general")
        files = body.get("files", [])  # 附件列表 [{name, mime_type, data(base64)}]
        web_search_enabled = body.get("web_search_enabled", False)
    except Exception:
        raise HTTPException(status_code=400, detail="无效的请求")

    # 构建用户消息：如果有附件，将文件信息编码进消息
    user_message = message
    if files:
        file_descriptions = []
        for f in files:
            file_descriptions.append(build_file_description(f))

        file_context = "\n\n".join(file_descriptions)
        if message:
            user_message = f"{message}\n\n--- 上传的文件 ---\n{file_context}"
        else:
            user_message = f"用户上传了以下文件：\n{file_context}\n请阅读文件内容并回复。"

    if not message.strip() and not files:
        raise HTTPException(status_code=400, detail="消息不能为空")

    def build_auto_title(msg: str, uploaded_files: list[dict]) -> str:
        title = " ".join((msg or "").strip().split())
        if not title and uploaded_files:
            first_name = uploaded_files[0].get("name", "附件")
            title = f"附件：{first_name}"
        if not title:
            return "新对话"
        return title[:28] + ("..." if len(title) > 28 else "")

    current_title = (session.get("title") or "").strip()
    if current_title in ("", "新对话"):
        auto_title = build_auto_title(display_message, files)
        if auto_title != "新对话":
            await update_session_title(token, auto_title, user["id"])

    history = await get_messages(token)
    system_prompt = get_prompt(scene)

    # 如果开启了联网搜索，在系统提示词中追加通知
    if web_search_enabled:
        system_prompt += "\n\n你拥有联网搜索工具（webSearchPrime）。当用户的问题涉及最新法规动态、时事新闻、你不确定的判例或需要最新信息时，请主动调用联网搜索工具获取最新资料，然后再结合法律知识回答。就像你使用法规检索和案例检索一样自然地使用它。"
    else:
        system_prompt += "\n\n注意：你当前没有联网搜索能力，仅基于自身训练数据和已提供的法规/案例检索结果回答。"

    # 合并工具
    from ai_client import LEGAL_TOOLS, get_ai_client, DEEPSEEK_MODEL
    from mcp_client import (
        get_mcp_tool_definitions,
        execute_mcp_tool,
        is_mcp_available,
        search_law_direct,
        search_case_direct,
    )

    mcp_tools = get_mcp_tool_definitions() if is_mcp_available() else []

    # 如果关闭了联网搜索，过滤掉智谱相关的 MCP 工具
    if mcp_tools and not web_search_enabled:
        mcp_tools = [
            t for t in mcp_tools
            if not t.get("function", {}).get("name", "").startswith("yuandian_zhipu_")
        ]
    all_tools = LEGAL_TOOLS + mcp_tools if mcp_tools else LEGAL_TOOLS

    def format_tool_name(tool_name: str) -> str:
        if tool_name == "search_law":
            return "法规检索"
        if tool_name == "search_case":
            return "案例检索"
        if tool_name == "get_law_detail":
            return "法条详情查询"
        if tool_name.startswith("yuandian_zhipu_"):
            return "联网搜索"
        if tool_name.startswith("yuandian_law_"):
            return "元典法规检索"
        if tool_name.startswith("yuandian_case_"):
            return "元典案例检索"
        if tool_name.startswith("yuandian_company_"):
            return "元典企业信息查询"
        if tool_name.startswith("yuandian_"):
            return "元典法律工具"
        if tool_name == "webSearchPrime" or tool_name.startswith("zhipu_"):
            return "联网搜索"
        return tool_name

    def summarize_tool_args(args: dict) -> str:
        for key in ("search_query", "query", "law_name", "article", "keyword", "keywords"):
            value = args.get(key)
            if value:
                text = str(value).strip()
                return text[:80] + ("..." if len(text) > 80 else "")
        return ""

    def sse_payload(payload: dict) -> str:
        return f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"

    def should_use_research_plan(scene_name: str, text: str) -> bool:
        if scene_name == "research":
            return True
        triggers = ("案例", "判例", "裁判", "类案", "案号", "法院", "再找", "检索")
        return any(word in text for word in triggers)

    def fallback_research_plan(text: str) -> list[dict]:
        query = " ".join((text or "").strip().split())[:80]
        if not query:
            return []
        if any(word in text for word in ("案例", "判例", "裁判", "类案", "案号", "法院")):
            return [{"type": "case", "query": query, "reason": "用户明确要求检索案例"}]
        return [
            {"type": "law", "query": query, "reason": "先检索相关法律依据"},
            {"type": "case", "query": query, "reason": "再检索相关裁判案例"},
        ]

    async def create_research_plan(client, text: str, hist: list[dict]) -> list[dict]:
        recent_context = "\n".join(
            f"{m.get('role')}: {(m.get('content') or '')[:300]}"
            for m in hist[-6:]
            if m.get("role") in ("user", "assistant") and m.get("content")
        )
        plan_prompt = f"""请为法律检索生成一个简洁检索计划。

要求：
1. 将用户问题拆成 2-4 个检索任务。
2. type 只能是 law 或 case。
3. query 使用适合法规/案例数据库的中文关键词，不要用完整长句。
4. 如果用户要求“再找案例/类似案例”，优先 case，并结合历史上下文补全关键词。
5. 只输出 JSON，不要输出解释。

历史上下文：
{recent_context or "无"}

当前问题：
{text}

JSON 格式：
{{"searches":[{{"type":"case","query":"房屋租赁合同纠纷 押金返还 提前退租","reason":"查找押金返还类案"}}]}}"""
        try:
            response = await client.chat.completions.create(
                model=DEEPSEEK_MODEL,
                max_tokens=800,
                temperature=0.2,
                messages=[
                    {"role": "system", "content": "你是法律检索策略助手，只输出可解析 JSON。"},
                    {"role": "user", "content": plan_prompt},
                ],
            )
            raw = response.choices[0].message.content or ""
            start = raw.find("{")
            end = raw.rfind("}")
            data = json.loads(raw[start:end + 1]) if start >= 0 and end >= start else {}
            searches = data.get("searches", [])
        except Exception:
            searches = fallback_research_plan(text)

        cleaned = []
        for item in searches:
            search_type = item.get("type")
            query = " ".join(str(item.get("query", "")).split())
            if search_type not in ("law", "case") or not query:
                continue
            cleaned.append({
                "type": search_type,
                "query": query[:120],
                "reason": str(item.get("reason", ""))[:120],
            })
            if len(cleaned) >= 4:
                break
        return cleaned or fallback_research_plan(text)

    def build_research_context(results: list[dict]) -> str:
        if not results:
            return ""
        parts = ["--- 自动检索计划与结果（供回答参考，不要逐字复述） ---"]
        for idx, item in enumerate(results, 1):
            kind = "案例检索" if item["type"] == "case" else "法规检索"
            parts.append(
                f"\n[{idx}] {kind}\n"
                f"检索词：{item['query']}\n"
                f"目的：{item.get('reason') or '补充法律依据'}\n"
                f"结果：\n{item.get('result') or '无结果'}"
            )
        parts.append(
            "--- 请基于以上可核验结果回答；如果结果不足，请明确说明。"
            "如果用户要求案例，请尽量列明案名、法院、案号、案由、裁判日期、核心事实、争议焦点、裁判观点和参考价值；"
            "如果用户要求法条，请尽量列明法规名称、条号、条文内容、效力层级/时效信息和适用关系。"
            "检索材料中没有显示的字段，不要编造，直接标注“检索结果未显示”。 ---"
        )
        return "\n".join(parts)

    def build_research_reference(item: dict, idx: int) -> dict | None:
        item_type = item.get("type")
        if item_type not in ("law", "case"):
            return None
        query = item.get("query") or ""
        title_prefix = "案例材料" if item_type == "case" else "法条材料"
        return {
            "id": f"{item_type}-{idx}-{abs(hash(query))}",
            "type": item_type,
            "title": f"{title_prefix}：{query[:60]}",
            "query": query,
            "reason": item.get("reason") or "",
            "content": str(item.get("raw_result") or item.get("result") or "")[:20000],
        }

    def build_research_references(results: list[dict]) -> list[dict]:
        references = []
        for idx, item in enumerate(results, 1):
            reference = build_research_reference(item, idx)
            if reference:
                references.append(reference)
        return references

    async def curate_research_references(client, text: str, results: list[dict]) -> list[dict]:
        raw_refs = build_research_references(results)
        if not raw_refs:
            return []

        source_text = "\n\n".join(
            f"[{idx}] type={ref['type']}\n"
            f"query={ref.get('query', '')}\n"
            f"reason={ref.get('reason', '')}\n"
            f"content={ref.get('content', '')[:7000]}"
            for idx, ref in enumerate(raw_refs, 1)
        )
        prompt = f"""请整理法律检索材料，输出可解析 JSON。

用户问题：
{text}

原始检索材料：
{source_text}

任务：
1. 从原始材料中抽取真实可展示的案例和法条，最多 10 条。
2. 对重复或高度相似材料去重。
3. 根据用户问题判断相关程度，给出 0-100 的 relevance_score。
4. 不要编造原始材料没有的信息；缺失字段填“检索结果未显示”。
5. 如果某条原始材料只是错误、空结果、工具提示或没有具体案例/法条内容，不要作为正式材料输出。

JSON 格式：
{{
  "references": [
    {{
      "type": "case",
      "title": "案名或可识别标题",
      "query": "对应检索词",
      "relevance_score": 86,
      "relevance_reason": "与用户所述事实和法律争点相关",
      "fields": {{
        "案名": "检索结果未显示",
        "法院": "检索结果未显示",
        "案号": "检索结果未显示",
        "案由": "检索结果未显示",
        "裁判日期": "检索结果未显示",
        "核心事实": "简述",
        "争议焦点": "简述",
        "裁判观点": "简述",
        "参考价值": "简述"
      }},
      "content": "基于原始材料整理出的完整展示文本"
    }},
    {{
      "type": "law",
      "title": "法规名称 + 条号",
      "query": "对应检索词",
      "relevance_score": 80,
      "relevance_reason": "说明相关性",
      "fields": {{
        "法规名称": "检索结果未显示",
        "条号": "检索结果未显示",
        "条文内容": "检索结果未显示",
        "效力/时效": "检索结果未显示",
        "适用关系": "简述"
      }},
      "content": "基于原始材料整理出的完整展示文本"
    }}
  ]
}}"""
        try:
            response = await client.chat.completions.create(
                model=DEEPSEEK_MODEL,
                max_tokens=4096,
                temperature=0.1,
                messages=[
                    {"role": "system", "content": "你是严谨的法律检索材料整理助手，只输出可解析 JSON。"},
                    {"role": "user", "content": prompt},
                ],
            )
            raw = response.choices[0].message.content or ""
            start = raw.find("{")
            end = raw.rfind("}")
            data = json.loads(raw[start:end + 1]) if start >= 0 and end >= start else {}
            curated = data.get("references") or []
        except Exception:
            curated = []

        cleaned = []
        for idx, ref in enumerate(curated, 1):
            ref_type = ref.get("type")
            if ref_type not in ("case", "law"):
                continue
            title = str(ref.get("title") or ("案例材料" if ref_type == "case" else "法条材料")).strip()
            content = str(ref.get("content") or "").strip()
            fields = ref.get("fields") if isinstance(ref.get("fields"), dict) else {}
            if not content and not fields:
                continue
            try:
                relevance_score = int(ref.get("relevance_score", 0))
            except Exception:
                relevance_score = 0
            cleaned.append({
                "id": f"curated-{ref_type}-{idx}",
                "type": ref_type,
                "title": title[:120],
                "query": str(ref.get("query") or "")[:160],
                "reason": str(ref.get("relevance_reason") or "模型整理后的检索材料")[:300],
                "relevance_score": max(0, min(100, relevance_score)),
                "fields": fields,
                "content": content[:12000],
                "curated": True,
            })

        return cleaned or raw_refs

    # 消息构建器
    def build_msgs(msg_text, hist, sys_prompt):
        msgs = []
        if sys_prompt:
            msgs.append({"role": "system", "content": sys_prompt})
        if hist:
            for m in hist[-60:]:
                out = {"role": m["role"]}
                if m.get("content"):
                    out["content"] = m["content"]
                if m.get("tool_calls"):
                    out["tool_calls"] = m["tool_calls"]
                if m.get("tool_call_id"):
                    out["tool_call_id"] = m["tool_call_id"]
                msgs.append(out)
        msgs.append({"role": "user", "content": msg_text})
        return msgs

    def find_web_search_tool_name() -> str:
        for tool in mcp_tools or []:
            name = tool.get("function", {}).get("name", "")
            if name.startswith("yuandian_zhipu_") or name == "webSearchPrime" or "web_search" in name.lower():
                return name
        return ""

    async def generate():
        try:
            client = get_ai_client()
            model_user_message = user_message
            if should_use_research_plan(scene, message):
                yield sse_payload({
                    "status": "正在拆解问题，生成检索计划...",
                    "tool": {"stage": "start", "name": "检索计划"},
                })
                plan = await create_research_plan(client, message or user_message, history)
                if plan:
                    yield sse_payload({
                        "status": f"已生成 {len(plan)} 个检索任务，开始执行...",
                        "tool": {"stage": "done", "name": "检索计划"},
                    })

                planned_results = []
                for idx, item in enumerate(plan, 1):
                    kind = "案例检索" if item["type"] == "case" else "法规检索"
                    yield sse_payload({
                        "status": f"正在执行第 {idx} 个任务：{kind} - {item['query']}",
                        "tool": {
                            "stage": "start",
                            "name": kind,
                            "query": item["query"],
                        },
                    })
                    if item["type"] == "case":
                        result = await search_case_direct(item["query"])
                    else:
                        result = await search_law_direct(item["query"])
                    result_text = str(result or "")
                    planned_results.append({
                        **item,
                        "result": result_text[:8000],
                        "raw_result": result_text,
                    })
                    reference = build_research_reference(planned_results[-1], idx)
                    if reference:
                        yield sse_payload({"references": [reference]})
                    yield sse_payload({
                        "status": f"第 {idx} 个检索任务完成，正在筛选结果...",
                        "tool": {"stage": "done", "name": kind},
                    })

                research_context = build_research_context(planned_results)
                yield sse_payload({
                    "status": "正在对检索材料去重、抽取要点并判断相关程度...",
                    "tool": {"stage": "start", "name": "材料整理"},
                })
                research_references = await curate_research_references(
                    client,
                    message or user_message,
                    planned_results,
                )
                if research_references:
                    yield sse_payload({
                        "references": research_references,
                        "replace_references": True,
                    })
                yield sse_payload({
                    "status": "检索材料整理完成，正在生成研究回答...",
                    "tool": {"stage": "done", "name": "材料整理"},
                })
                if research_context:
                    model_user_message = f"{user_message}\n\n{research_context}"

            if web_search_enabled and is_mcp_available():
                web_tool_name = find_web_search_tool_name()
                original_web_query = (message or display_message or user_message).strip()
                if web_tool_name and original_web_query:
                    web_query = original_web_query
                    try:
                        rewrite_response = await client.chat.completions.create(
                            model=DEEPSEEK_MODEL,
                            max_tokens=120,
                            temperature=0.2,
                            messages=[
                                {
                                    "role": "system",
                                    "content": (
                                        "你是联网搜索查询改写器。请把用户问题改写成适合搜索引擎和新闻/资料检索的中文查询词。"
                                        "保留关键主体、时间、地域、法律概念和争议焦点；去掉寒暄和无关指令。"
                                        "只输出一行查询词，不要解释，不要加引号。"
                                    ),
                                },
                                {"role": "user", "content": original_web_query[:1200]},
                            ],
                        )
                        await record_response_usage(
                            user=user,
                            feature="web_search_rewrite",
                            model=DEEPSEEK_MODEL,
                            response=rewrite_response,
                        )
                        rewritten = (rewrite_response.choices[0].message.content or "").strip()
                        rewritten = rewritten.strip("`\"'“”‘’ \n")
                        if rewritten:
                            web_query = rewritten[:500]
                    except Exception as rewrite_error:
                        print(f"[WebSearch] 查询改写失败，使用原始问题: {rewrite_error}")

                    readable_tool_name = format_tool_name(web_tool_name)
                    short_query = web_query[:80] + ("..." if len(web_query) > 80 else "")
                    yield sse_payload({
                        "status": f"正在调用{readable_tool_name}: {short_query}",
                        "tool": {
                            "stage": "start",
                            "name": readable_tool_name,
                            "query": short_query,
                        },
                    })
                    web_args = {
                        "search_query": web_query,
                        "count": 5,
                    }
                    web_result = await execute_mcp_tool(web_tool_name, web_args)
                    yield sse_payload({
                        "status": f"{readable_tool_name}完成，正在结合最新材料...",
                        "tool": {
                            "stage": "done",
                            "name": readable_tool_name,
                        },
                    })
                    yield sse_payload({
                        "references": [{
                            "id": f"web-forced-{uuid.uuid4().hex}",
                            "type": "web",
                            "title": f"联网搜索: {short_query}",
                            "query": short_query,
                            "reason": "用户已开启联网搜索，系统先改写问题并获取最新网络材料。",
                            "content": str(web_result or "")[:20000],
                        }]
                    })
                    web_context = (
                        "\n\n--- 联网搜索材料（用户已开启联网搜索，回答时请优先结合这些最新材料，并说明检索材料的限制）---\n"
                        f"原始问题：{original_web_query}\n"
                        f"改写检索词：{web_query}\n"
                        f"{str(web_result or '')[:20000]}"
                    )
                    model_user_message = f"{model_user_message}{web_context}"
                else:
                    yield sse_payload({
                        "status": "联网搜索已开启，但当前环境未加载到 webSearchPrime 工具，将继续基于其他材料回答。",
                        "tool": {
                            "stage": "failed",
                            "name": "联网搜索",
                        },
                    })

            messages = build_msgs(model_user_message, history, system_prompt)
            new_messages_start = len(messages) - 1

            max_turns = 5
            full_response = ""

            for turn in range(max_turns):
                if turn == 0:
                    yield sse_payload({"status": "正在理解问题，判断是否需要检索资料..."})
                else:
                    yield sse_payload({"status": "正在结合检索结果继续分析..."})

                kwargs = {
                    "model": DEEPSEEK_MODEL,
                    "max_tokens": 4096,
                    "messages": messages,
                    "stream": True,
                    "stream_options": {"include_usage": True},
                    "temperature": 0.7,
                }
                if all_tools:
                    kwargs["tools"] = all_tools

                try:
                    stream = await client.chat.completions.create(**kwargs)
                except Exception as stream_error:
                    if "stream_options" not in str(stream_error):
                        raise
                    kwargs.pop("stream_options", None)
                    stream = await client.chat.completions.create(**kwargs)

                content_parts = []
                tool_calls_map = {}
                turn_usage = None

                async for chunk in stream:
                    if getattr(chunk, "usage", None):
                        turn_usage = extract_usage(chunk.usage)
                    if not chunk.choices:
                        continue
                    delta = chunk.choices[0].delta

                    if delta.content:
                        content_parts.append(delta.content)
                        full_response += delta.content
                        yield sse_payload({"content": delta.content})

                    if delta.tool_calls:
                        for tc in delta.tool_calls:
                            idx = tc.index
                            if idx not in tool_calls_map:
                                tool_calls_map[idx] = {
                                    "id": tc.id or "",
                                    "function": {"name": "", "arguments": ""},
                                }
                            if tc.id:
                                tool_calls_map[idx]["id"] = tc.id
                            if tc.function:
                                if tc.function.name:
                                    tool_calls_map[idx]["function"]["name"] += tc.function.name
                                if tc.function.arguments:
                                    tool_calls_map[idx]["function"]["arguments"] += tc.function.arguments

                await record_usage(user=user, feature="chat", model=DEEPSEEK_MODEL, usage=turn_usage)

                final_content = "".join(content_parts)
                if not tool_calls_map:
                    if final_content:
                        messages.append({
                            "role": "assistant",
                            "content": final_content,
                        })
                    break

                assistant_msg = {
                    "role": "assistant",
                    "content": final_content or None,
                }
                tool_calls_list = []
                for idx in sorted(tool_calls_map.keys()):
                    tc = tool_calls_map[idx]
                    tc["type"] = "function"
                    if not tc["function"]["arguments"]:
                        tc["function"]["arguments"] = "{}"
                    try:
                        json.loads(tc["function"]["arguments"])
                    except json.JSONDecodeError:
                        tc["function"]["arguments"] = "{}"
                    tool_calls_list.append(tc)
                assistant_msg["tool_calls"] = tool_calls_list
                messages.append(assistant_msg)

                for tc in tool_calls_list:
                    tool_name = tc["function"]["name"]
                    try:
                        args = json.loads(tc["function"]["arguments"])
                    except json.JSONDecodeError:
                        args = {}

                    readable_tool_name = format_tool_name(tool_name)
                    readable_args = summarize_tool_args(args)
                    yield sse_payload({
                        "status": f"正在调用{readable_tool_name}..."
                        if not readable_args else f"正在调用{readable_tool_name}：{readable_args}",
                        "tool": {
                            "stage": "start",
                            "name": readable_tool_name,
                            "query": readable_args,
                        },
                    })

                    if tool_name.startswith("yuandian_"):
                        result = await execute_mcp_tool(tool_name, args)
                    elif is_mcp_available() and tool_name == "search_law":
                        result = await search_law_direct(args.get("query", ""))
                    elif is_mcp_available() and tool_name == "search_case":
                        result = await search_case_direct(args.get("query", ""))
                    elif is_mcp_available() and tool_name == "get_law_detail":
                        result = await search_law_direct(
                            f"{args.get('law_name', '')} {args.get('article', '')}"
                        )
                    else:
                        result = json.dumps({
                            "info": "元典 MCP 未启用，AI 将基于训练数据中的法律知识回答"
                        }, ensure_ascii=False)

                    yield sse_payload({
                        "status": f"{readable_tool_name}完成，正在整理结果...",
                        "tool": {
                            "stage": "done",
                            "name": readable_tool_name,
                        },
                    })
                    if tool_name == "search_law" or tool_name == "get_law_detail" or tool_name.startswith("yuandian_law_"):
                        yield sse_payload({
                            "references": [{
                                "id": f"law-tool-{tc['id']}",
                                "type": "law",
                                "title": f"法条材料：{readable_args or readable_tool_name}",
                                "query": readable_args,
                                "reason": "模型主动调用法规/法条检索",
                                "content": str(result or "")[:20000],
                            }]
                        })
                    elif tool_name == "search_case" or tool_name.startswith("yuandian_case_"):
                        yield sse_payload({
                            "references": [{
                                "id": f"case-tool-{tc['id']}",
                                "type": "case",
                                "title": f"案例材料：{readable_args or readable_tool_name}",
                                "query": readable_args,
                                "reason": "模型主动调用案例检索",
                                "content": str(result or "")[:20000],
                            }]
                        })
                    elif tool_name.startswith("yuandian_zhipu_"):
                        yield sse_payload({
                            "references": [{
                                "id": f"web-tool-{tc['id']}",
                                "type": "web",
                                "title": f"联网搜索：{readable_args or readable_tool_name}",
                                "query": readable_args,
                                "reason": "模型主动调用联网搜索",
                                "content": str(result or "")[:20000],
                            }]
                        })

                    messages.append({
                        "role": "tool",
                        "tool_call_id": tc["id"],
                        "content": result,
                    })

                yield sse_payload({"content": "\n\n"})

            # 只保存本轮新增消息，避免历史消息重复入库；检索上下文不写入用户消息。
            messages_to_save = [dict(m) for m in messages[new_messages_start:]]
            for msg in messages_to_save:
                if msg.get("role") == "user":
                    msg["content"] = display_message if display_message.strip() else user_message
                    break
            await save_new_messages(token, messages_to_save, 0)
            yield "data: [DONE]\n\n"
        except Exception as e:
            yield sse_payload({"error": str(e)})

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@app.post("/api/research/followup")
async def research_followup(request: Request):
    """围绕当前检索工作台结果继续追问，不触碰普通对话历史。"""
    from ai_client import get_ai_client, DEEPSEEK_MODEL

    user = await require_current_user(request)
    try:
        body = await request.json()
        question = (body.get("question") or "").strip()
        original_query = (body.get("query") or "").strip()
        answer = (body.get("answer") or "").strip()
        references = body.get("references") or []
        files = body.get("files", []) or []
        record_id = (body.get("record_id") or "").strip()
        conversation = body.get("conversation") or []
    except Exception:
        raise HTTPException(status_code=400, detail="无效的请求")

    if not question:
        raise HTTPException(status_code=400, detail="追问不能为空")

    def sse_payload(payload: dict) -> str:
        return f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"

    ref_context = json.dumps(references[:20], ensure_ascii=False)[:30000]
    conversation_context = json.dumps(conversation[-12:], ensure_ascii=False)[:16000] if isinstance(conversation, list) else ""
    file_context = ""
    if files:
        material_parts = []
        for index, file_info in enumerate(files[:8], 1):
            try:
                material_parts.append(f"【追问附件{index}】\n{build_file_description(file_info, max_chars=10000)}")
            except Exception as file_error:
                filename = file_info.get("name", f"附件{index}") if isinstance(file_info, dict) else f"附件{index}"
                material_parts.append(f"【追问附件{index}: {filename}】\n（读取失败：{file_error}）")
        file_context = "\n\n".join(material_parts)

    async def generate():
        answer_text = ""
        try:
            client = get_ai_client()
            prompt = f"""请基于当前法规/案例检索工作台的上下文回答用户追问。

原始研究问题：
{original_query or "未提供"}

当前研究回答：
{answer[:12000] or "未提供"}

此前追问对话：
{conversation_context or "未提供"}

当前检索材料：
{ref_context or "未提供"}

用户追问：
{question}

用户本次追问上传的附件：
{file_context or "无"}

要求：
1. 只基于当前上下文和检索材料回答，不要声称重新检索。
2. 如果现有材料不足以回答，请明确说明还需要补充检索什么。
3. 涉及案例或法条时，引用右侧材料中的案号、案名、法条或链接字段。"""
            stream_kwargs = {
                "model": DEEPSEEK_MODEL,
                "max_tokens": 3000,
                "temperature": 0.35,
                "stream": True,
                "stream_options": {"include_usage": True},
                "messages": [
                    {"role": "system", "content": "你是严谨的法律研究追问助手，请用中文回答。"},
                    {"role": "user", "content": prompt},
                ],
            }
            try:
                stream = await client.chat.completions.create(**stream_kwargs)
            except Exception as stream_error:
                if "stream_options" not in str(stream_error):
                    raise
                stream_kwargs.pop("stream_options", None)
                stream = await client.chat.completions.create(**stream_kwargs)
            answer_usage = None
            async for chunk in stream:
                if getattr(chunk, "usage", None):
                    answer_usage = extract_usage(chunk.usage)
                if not chunk.choices:
                    continue
                delta = chunk.choices[0].delta
                if delta.content:
                    answer_text += delta.content
                    yield sse_payload({"content": delta.content})
            await record_usage(user=user, feature="research_followup", model=DEEPSEEK_MODEL, usage=answer_usage)
            if record_id:
                from db import add_research_message
                await add_research_message(
                    record_id=record_id,
                    user_id=user["id"],
                    role="user",
                    content=question,
                    files=summarize_uploaded_files(files),
                )
                await add_research_message(
                    record_id=record_id,
                    user_id=user["id"],
                    role="assistant",
                    content=answer_text,
                    files=[],
                )
            yield "data: [DONE]\n\n"
        except Exception as e:
            yield sse_payload({"error": str(e)})

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@app.post("/api/chat-clear")
async def chat_clear(request: Request):
    from db import get_session, clear_messages
    user = await require_current_user(request)
    token = request.cookies.get("session_token")
    session = await get_session(token, user["id"]) if token else None
    if not session:
        raise HTTPException(status_code=401, detail="请先登录")
    await clear_messages(token)
    return {"success": True}


@app.get("/api/research/records")
async def research_records(request: Request):
    from db import list_research_records
    user = await require_current_user(request)
    return {"records": await list_research_records(user["id"])}


@app.get("/api/research/records/{record_id}")
async def research_record_detail(record_id: str, request: Request):
    from db import get_research_record
    user = await require_current_user(request)
    record = await get_research_record(record_id, user["id"])
    if not record:
        raise HTTPException(status_code=404, detail="检索记录不存在")
    return {"record": record}


@app.post("/api/research/records/delete")
async def research_record_delete(request: Request):
    from db import delete_research_record
    user = await require_current_user(request)
    try:
        body = await request.json()
        record_id = (body.get("record_id") or "").strip()
    except Exception:
        raise HTTPException(status_code=400, detail="无效的请求")
    if not record_id:
        raise HTTPException(status_code=400, detail="记录 ID 不能为空")
    deleted = await delete_research_record(record_id, user["id"])
    if not deleted:
        raise HTTPException(status_code=404, detail="检索记录不存在")
    return {"success": True}


@app.delete("/api/research/records/{record_id}")
async def research_record_delete_by_id(record_id: str, request: Request):
    from db import delete_research_record
    user = await require_current_user(request)
    deleted = await delete_research_record(record_id, user["id"])
    if not deleted:
        raise HTTPException(status_code=404, detail="检索记录不存在")
    return {"success": True}


@app.post("/api/research/stream")
async def research_stream(request: Request):
    """法规/案例检索工作台专用接口：独立检索、整理、回答、保存研究记录。"""
    from ai_client import get_ai_client, DEEPSEEK_MODEL, RESEARCH_PRO_MODEL
    from mcp_client import (
        get_case_details_direct,
        is_mcp_available,
        search_case_authoritative_direct,
        search_case_ordinary_direct,
        search_case_vector_direct,
        search_law_direct,
    )
    from db import save_research_record

    user = await require_current_user(request)
    try:
        body = await request.json()
        query_text = (body.get("query") or "").strip()
        need_cases = bool(body.get("need_cases", True))
        need_laws = bool(body.get("need_laws", False))
        region = (body.get("region") or "").strip()
        date_start = (body.get("date_start") or "").strip()
        date_end = (body.get("date_end") or "").strip()
        focus = (body.get("focus") or "类案和问题回答并重").strip()
        filter_strength = (body.get("filter_strength") or "standard").strip()
        files = body.get("files", []) or []
    except Exception:
        raise HTTPException(status_code=400, detail="无效的请求")

    if not query_text:
        raise HTTPException(status_code=400, detail="检索需求不能为空")

    display_query_text = query_text
    file_context = ""
    if files:
        material_parts = []
        for index, file_info in enumerate(files[:8], 1):
            try:
                material_parts.append(f"【上传材料{index}】\n{build_file_description(file_info, max_chars=12000)}")
            except Exception as file_error:
                filename = file_info.get("name", f"材料{index}") if isinstance(file_info, dict) else f"材料{index}"
                material_parts.append(f"【上传材料{index}: {filename}】\n（读取失败：{file_error}）")
        file_context = "\n\n".join(material_parts)
        query_text = f"{query_text}\n\n--- 用户上传的研究材料 ---\n{file_context}"

    strength_config = {
        "loose": {"name": "宽松", "threshold": 35, "limit": 15, "instruction": "宁可多保留弱相关候选，方便人工复核。"},
        "standard": {"name": "标准", "threshold": 55, "limit": 10, "instruction": "兼顾相关度和材料数量，保留中高相关材料。"},
        "strict": {"name": "严格", "threshold": 75, "limit": 6, "instruction": "只保留事实和法律问题高度相近的强相关材料。"},
    }
    strength = strength_config.get(filter_strength, strength_config["standard"])
    min_keep_score = strength["threshold"] if filter_strength == "strict" else (30 if filter_strength == "standard" else 0)

    def sse_payload(payload: dict) -> str:
        return f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"

    def parse_json_object(raw: str) -> dict:
        start = raw.find("{")
        end = raw.rfind("}")
        if start >= 0 and end >= start:
            return json.loads(raw[start:end + 1])
        return {}

    def build_keyword_query(source: str, item_type: str = "case", purpose: str = "") -> str:
        text = re.sub(r"--- 用户上传的研究材料 ---.*", "", source or "", flags=re.S)
        text = re.sub(r"[0-9]{4}年|[0-9]+月|[0-9]+日|共计|四个月份|背景[:：]|麻烦你|检索一下|需要|确认|是否|等等", " ", text)
        text = re.sub(r"[，。；;、（）()《》“”\"'：:,.!?？!\n\r\t]", " ", text)
        # 兜底词只从当前问题动态提取，不预置任何具体业务、案由或请求权。
        # 主流程仍由检索规划模型按不同案例库生成更合适的查询。
        text = re.sub(r"(?:以及|并且|但是|仍然|其中|例如|基于|能否|是否|应当|已经|相关|前述|等等|要求|需要)", " ", text)
        raw_tokens = re.findall(r"[\u4e00-\u9fa5A-Za-z0-9]{2,16}", text)
        tokens = []
        if region:
            tokens.append(region)
        for token in raw_tokens:
            token = token.strip()
            if token and token not in tokens:
                tokens.append(token)
            if len(tokens) >= 10:
                break
        defaults = [] if item_type == "case" else ["法律依据"]
        for token in defaults:
            if token not in tokens:
                tokens.append(token)
        return " ".join(tokens[:12]).strip()[:90]

    def fallback_plan() -> list[dict]:
        case_plan = []
        law_plan = []
        if need_cases:
            base = build_keyword_query(query_text, "case")
            broad = " ".join(base.split()[:6]) or display_query_text[:60]
            region_word = region or ("北京" if "北京" in query_text else "")
            ordinary_base = " ".join(token for token in (region_word, broad) if token).strip()
            core_tokens = [token for token in broad.split() if token and token != region_word]

            def authority_query(offset: int, suffix: str) -> str:
                if not core_tokens:
                    return suffix
                rotated = core_tokens[offset:] + core_tokens[:offset]
                return " ".join(rotated[:2] + [suffix])[:60]

            case_plan = [
                {"type": "case", "channel": "semantic", "query": display_query_text[:500], "reason": "按完整事实结构召回相似案例"},
                {"type": "case", "channel": "semantic", "query": f"{display_query_text[:420]}，重点识别法律关系、案由和请求权基础", "reason": "按法律关系和请求权基础召回相似案例"},
                {"type": "case", "channel": "semantic", "query": f"{display_query_text[:420]}，重点关注责任成立的构成要件和适用条件", "reason": "按构成要件和责任条件召回相似案例"},
                {"type": "case", "channel": "semantic", "query": f"{display_query_text[:420]}，重点关注相对方抗辩、举证责任和裁判结果", "reason": "按抗辩、举证和裁判结果召回相似案例"},
                {"type": "case", "channel": "ordinary", "query": ordinary_base[:90], "arguments": {"qw": ordinary_base[:90], "search_mode": "or", "top_k": 15}, "reason": "普通案例库宽泛检索核心事实"},
                {"type": "case", "channel": "ordinary", "query": f"{ordinary_base} 案由 争议焦点"[:90], "arguments": {"qw": f"{ordinary_base} 案由 争议焦点"[:90], "search_mode": "or", "top_k": 15}, "reason": "普通案例库检索案由和争议焦点"},
                {"type": "case", "channel": "ordinary", "query": f"{ordinary_base} 构成要件 责任认定"[:90], "arguments": {"qw": f"{ordinary_base} 构成要件 责任认定"[:90], "search_mode": "or", "top_k": 15}, "reason": "普通案例库检索构成要件和责任认定"},
                {"type": "case", "channel": "ordinary", "query": f"{ordinary_base} 请求权 抗辩 举证责任"[:90], "arguments": {"qw": f"{ordinary_base} 请求权 抗辩 举证责任"[:90], "search_mode": "or", "top_k": 15}, "reason": "普通案例库检索请求权、抗辩和举证规则"},
                {"type": "case", "channel": "authoritative", "query": authority_query(0, "法律关系"), "arguments": {"qw": authority_query(0, "法律关系"), "search_mode": "or", "top_k": 15}, "reason": "权威案例库检索法律关系裁判规则"},
                {"type": "case", "channel": "authoritative", "query": authority_query(1, "责任认定"), "arguments": {"qw": authority_query(1, "责任认定"), "search_mode": "or", "top_k": 15}, "reason": "权威案例库检索责任认定规则"},
                {"type": "case", "channel": "authoritative", "query": authority_query(2, "构成要件"), "arguments": {"qw": authority_query(2, "构成要件"), "search_mode": "or", "top_k": 15}, "reason": "权威案例库检索构成要件规则"},
                {"type": "case", "channel": "authoritative", "query": authority_query(3, "请求权抗辩"), "arguments": {"qw": authority_query(3, "请求权抗辩"), "search_mode": "or", "top_k": 15}, "reason": "权威案例库检索请求权与抗辩规则"},
            ]
        if need_laws:
            base = build_keyword_query(query_text, "law")
            law_plan = [
                {"type": "law", "query": f"{base} 权利义务 相关法条", "reason": "检索基础权利义务规则"},
                {"type": "law", "query": f"{base} 责任承担 法律依据", "reason": "检索责任承担规则"},
                {"type": "law", "query": f"{base} 构成要件 适用条件", "reason": "检索请求权构成和适用条件"},
                {"type": "law", "query": f"{base} 司法解释 效力时效", "reason": "检索相关司法解释和现行效力"},
            ]

        if need_cases and need_laws:
            return case_plan[:12] + law_plan[:4]
        if need_cases:
            return case_plan[:12]
        if need_laws:
            return law_plan[:4]
        return [{"type": "case", "query": build_keyword_query(query_text, "case"), "reason": "检索相关案例"}]

    async def create_plan(client) -> list[dict]:
        type_requirement = (
            "必须恰好生成 16 个任务：12 个 case（semantic、ordinary、authoritative 各 4 个）、4 个 law。"
            if need_cases and need_laws else
            "必须恰好生成 12 个 case 任务：semantic、ordinary、authoritative 各 4 个。"
            if need_cases else
            "必须恰好生成 4 个 law 任务。"
        )
        prompt = f"""请为法规/案例检索生成检索任务，只输出 JSON。
用户问题：{query_text}
地域限制：{region or "无"}
需要案例：{need_cases}
需要法条：{need_laws}
输出侧重：{focus}
筛选强度：{strength["name"]}。{strength["instruction"]}

要求：
1. type 只能是 case 或 law。
2. {type_requirement}
3. 每个 case 任务只属于一个 channel，channel 只能是 semantic、ordinary、authoritative。三个通道必须各自独立规划 4 个任务，不得把同一个任务机械复制到三路：
   - semantic：query 使用完整、简洁的争议事实描述，四个任务分别覆盖相似事实、案由、返还条件、请求权与抗辩；
   - ordinary：query 使用 4-8 个适合普通裁判文书库的关键词；同时提供 arguments，仅含 qw、search_mode、top_k；
   - authoritative：query 使用 2-4 个裁判规则关键词，侧重典型/参考案例可能采用的规范表达；同时提供 arguments，仅含 qw、search_mode、top_k，使用 or 宽松召回。
4. law 任务围绕法律关系、权利义务、构成要件、责任承担和司法解释设计，彼此不得只是同义改写。
5. ordinary、authoritative 和 law 的 query 必须是适合法律数据库的中文关键词组合，不要复制用户背景长句；semantic 按第 3 条使用精炼的事实描述。
6. 如果需要北京等地域案例，把地域写进 case query；法规检索一般不要加入地域，地方性法规除外。
7. 除 semantic 外，query 最长不超过 60 个中文字符。

JSON：{{"searches":[{{"type":"case","channel":"semantic","query":"买方已付款但卖方逾期交付，买方请求承担违约责任","reason":"按相似事实检索"}},{{"type":"case","channel":"ordinary","query":"买卖合同 逾期交付 违约责任","arguments":{{"qw":"买卖合同 逾期交付 违约责任","search_mode":"or","top_k":15}},"reason":"普通库关键词检索"}},{{"type":"case","channel":"authoritative","query":"逾期交付 违约责任","arguments":{{"qw":"逾期交付 违约责任","search_mode":"or","top_k":15}},"reason":"权威库裁判规则检索"}},{{"type":"law","query":"买卖合同 逾期交付 违约责任 法律依据","reason":"检索相关法条"}}]}}"""
        try:
            response = await client.chat.completions.create(
                model=RESEARCH_PRO_MODEL,
                max_tokens=3000,
                temperature=0.15,
                messages=[
                    {"role": "system", "content": "你是法律检索策略助手，只输出可解析 JSON。"},
                    {"role": "user", "content": prompt},
                ],
            )
            searches = parse_json_object(response.choices[0].message.content or "").get("searches") or []
        except Exception:
            searches = []

        cleaned = []
        for item in searches:
            item_type = item.get("type")
            item_channel = str(item.get("channel") or "").strip() if item_type == "case" else ""
            item_query = " ".join(str(item.get("query") or "").split())
            if item_type not in ("case", "law") or not item_query:
                continue
            if (
                (item_type != "case" or item_channel != "semantic")
                and (len(item_query) > 70 or any(marker in item_query for marker in ("背景", "麻烦", "2023年", "共计", "但仍收取")))
            ):
                item_query = build_keyword_query(f"{item_query} {query_text}", item_type, str(item.get("reason") or ""))
            if item_type == "case" and not need_cases:
                continue
            if item_type == "law" and not need_laws:
                continue
            candidate = {"type": item_type, "query": item_query[:90], "reason": str(item.get("reason") or "")[:160]}
            if item_type == "case":
                channel = item_channel
                if channel not in ("semantic", "ordinary", "authoritative"):
                    continue

                def clean_case_args(value: dict, *, word_limit: int) -> dict:
                    value = value if isinstance(value, dict) else {}
                    allowed = ("qw", "search_mode", "top_k")
                    cleaned_args = {key: value.get(key) for key in allowed if value.get(key) not in (None, "", [], {})}
                    if cleaned_args.get("qw"):
                        cleaned_args["qw"] = " ".join(str(cleaned_args["qw"]).split()[:word_limit])
                    cleaned_args["search_mode"] = "or"
                    cleaned_args["top_k"] = 15
                    return cleaned_args

                candidate["channel"] = channel
                if channel != "semantic":
                    word_limit = 8 if channel == "ordinary" else 4
                    arguments = clean_case_args(item.get("arguments") or {}, word_limit=word_limit)
                    arguments.setdefault("qw", " ".join(item_query.split()[:word_limit]))
                    candidate["arguments"] = arguments
                else:
                    candidate["query"] = " ".join(str(item.get("query") or "").split())[:500]
            if not any(existing["type"] == candidate["type"] and existing.get("channel") == candidate.get("channel") and existing["query"] == candidate["query"] for existing in cleaned):
                cleaned.append(candidate)
            if len(cleaned) >= 20:
                break

        # 案例三路各自独立补齐 4 个任务；法规单独补齐 4 个任务。
        fallback_items = fallback_plan()
        balanced = []
        if need_cases or not need_laws:
            for channel in ("semantic", "ordinary", "authoritative"):
                typed = [item for item in cleaned if item["type"] == "case" and item.get("channel") == channel]
                typed.extend(
                    item for item in fallback_items
                    if item["type"] == "case" and item.get("channel") == channel
                    and not any(existing["query"] == item["query"] for existing in typed)
                )
                balanced.extend(typed[:4])
        if need_laws:
            typed = [item for item in cleaned if item["type"] == "law"]
            typed.extend(
                item for item in fallback_items
                if item["type"] == "law" and not any(existing["query"] == item["query"] for existing in typed)
            )
            balanced.extend(typed[:4])
        return balanced or fallback_items

    def normalize_reference(ref: dict, idx: int, prefix: str = "research") -> dict | None:
        ref_type = ref.get("type")
        if ref_type not in ("case", "law"):
            return None
        fields = ref.get("fields") if isinstance(ref.get("fields"), dict) else {}
        content = str(ref.get("content") or "").strip()
        if not fields and not content:
            return None
        try:
            score = int(ref.get("relevance_score", 0))
        except Exception:
            score = 0

        normalized_fields = {}
        expected = (
            ["scid", "案例库", "来源类型", "案名", "法院", "案号", "案由", "裁判日期", "审判程序", "案件类别", "文书种类", "省份", "核心事实", "争议焦点", "裁判观点", "裁判结果", "援引法条", "参考价值", "链接"]
            if ref_type == "case"
            else ["法规名称", "条号", "条文内容", "效力/时效", "适用关系", "链接"]
        )
        for key in expected:
            value = fields.get(key) or fields.get(key.lower()) or "未提取到信息"
            normalized_fields[key] = str(value).strip() or "未提取到信息"

        link = str(ref.get("link") or normalized_fields.get("链接") or "").strip()
        if not link:
            link = "未提取到信息"
        normalized_fields["链接"] = link

        return {
            "id": f"{prefix}-{ref_type}-{idx}",
            "source_id": str(ref.get("source_id") or ref.get("id") or f"{prefix}-{ref_type}-{idx}"),
            "type": ref_type,
            "title": str(ref.get("title") or normalized_fields.get("案名") or normalized_fields.get("法规名称") or ("案例材料" if ref_type == "case" else "法条材料"))[:140],
            "query": str(ref.get("query") or "")[:180],
            "reason": str(ref.get("relevance_reason") or ref.get("reason") or "模型筛选后的检索材料")[:300],
            "relevance_score": max(0, min(100, score)),
            "link": link,
            "fields": normalized_fields,
            "content": content[:12000],
            "curated": True,
            "case_library": str(ref.get("case_library") or "") if ref_type == "case" else "",
            "authority_type": str(ref.get("authority_type") or "") if ref_type == "case" else "",
            "matched_by": list(ref.get("matched_by") or []) if ref_type == "case" else [],
            "scid": str(ref.get("scid") or fields.get("scid") or "") if ref_type == "case" else "",
            "ordinary_id": str(ref.get("ordinary_id") or "") if ref_type == "case" else "",
            "authority_id": str(ref.get("authority_id") or "") if ref_type == "case" else "",
        }

    async def extract_candidate_references(client, item: dict, idx: int) -> list[dict]:
        raw_result = str(item.get("result") or "")
        if not result_has_material(raw_result):
            return []
        ref_type = item.get("type")
        prompt = f"""请从单次 MCP 检索结果中抽取候选法律材料，只输出 JSON。
用户问题：{query_text}
地域限制：{region or "无"}
检索类型：{ref_type}
检索词：{item.get("query") or ""}
检索目的：{item.get("reason") or ""}

MCP 原始返回：
{raw_result[:12000]}

要求：
1. 从原始返回中抽取候选案例/法条，最多 8 条。
2. 不要因为字段不完整而丢弃材料；缺失字段统一写“未提取到信息”。
3. 必须尝试提取链接/详情页/来源 URL；没有则写“未提取到信息”。
4. 只基于原文，不要编造。

JSON：
{{"references":[{{"type":"case","title":"案例标题","query":"{item.get("query") or ""}","link":"未提取到信息","relevance_score":60,"relevance_reason":"候选相关性说明","fields":{{"案名":"未提取到信息","法院":"未提取到信息","案号":"未提取到信息","案由":"未提取到信息","裁判日期":"未提取到信息","核心事实":"未提取到信息","争议焦点":"未提取到信息","裁判观点":"未提取到信息","参考价值":"未提取到信息","链接":"未提取到信息"}},"content":"候选材料完整摘要"}}]}}"""
        try:
            response = await client.chat.completions.create(
                model=DEEPSEEK_MODEL,
                max_tokens=4096,
                temperature=0.1,
                messages=[
                    {"role": "system", "content": "你是法律材料抽取助手，只输出可解析 JSON。"},
                    {"role": "user", "content": prompt},
                ],
            )
            refs = parse_json_object(response.choices[0].message.content or "").get("references") or []
        except Exception:
            refs = []

        cleaned = []
        for ref_index, ref in enumerate(refs, 1):
            if ref_type in ("case", "law"):
                ref["type"] = ref.get("type") or ref_type
            ref["query"] = ref.get("query") or item.get("query") or ""
            ref["source_id"] = ref.get("source_id") or f"candidate-{idx}-{ref_index}"
            normalized = normalize_reference(ref, ref_index, prefix=f"candidate-{idx}")
            if normalized:
                cleaned.append(normalized)
        return cleaned

    async def select_candidate_scids(client, item: dict, idx: int) -> list[dict]:
        raw_result = str(item.get("result") or "")
        if not result_has_material(raw_result):
            return []
        search_channel = str(item.get("channel") or "semantic")
        channel_label = {"semantic": "语义检索", "ordinary": "普通案例库", "authoritative": "权威案例库"}.get(search_channel, search_channel)
        prompt = f"""请从 MCP 案例检索结果中筛选与用户问题匹配的案例，只输出 JSON。
用户问题：{query_text}
地域限制：{region or "无"}
检索词：{item.get("query") or ""}
检索目的：{item.get("reason") or ""}
筛选强度：{strength["name"]}，相关度阈值约 {strength["threshold"]}。{strength["instruction"]}

评分口径：
- 90-100：事实、地域、案由/法律问题高度一致，可直接强支持。
- 80-89：法律争点或裁判规则高度一致，事实有差异但可直接参考。
- 70-79：请求权基础、裁判规则或抗辩逻辑相关，可作为类案参考。
- 60-69：背景规则相关，需要谨慎类比。
- 50-59：弱相关，仅作补充线索。
- 50 以下：通常不展示，除非宽松模式下有检索线索价值。
不要因为地区、金额、主体类型、字段缺失而机械降分；只要法律争点或裁判规则可用于回答用户问题，应给 70 分以上。

        MCP {channel_label}原始返回：
        {raw_result[:36000]}

任务：
1. 只筛选与用户问题相关的案例。
        2. 保留各渠道的原始标识：语义检索的 scid、普通案例的 id 写入 ordinary_id、权威案例的 id 写入 authority_id。
        3. 本次结果固定来自 {search_channel}，matched_by 必须包含 {search_channel}。
        4. case_library 填 {search_channel}；权威案例同时提取 authority_type（如典型案例、参考案例）。
        5. 同时提取案名/标题、案号、法院、链接或详情链接；缺失写“未提取到信息”。
        6. 给出 relevance_score 0-100 和筛选理由。
        7. relevance_score 低于 {strength["threshold"]} 的通常不要输出，除非具有明显参考价值。
        8. 最多输出 {strength["limit"]} 条。

JSON：
        {{"selected":[{{"scid":"语义检索 scid 或未提取到信息","ordinary_id":"普通案例 id 或未提取到信息","authority_id":"权威案例 id 或未提取到信息","case_no":"案号或未提取到信息","case_library":"mixed","authority_type":"典型案例或未提取到信息","matched_by":["semantic","ordinary","authoritative"],"title":"案例标题","court":"法院或未提取到信息","link":"链接或未提取到信息","relevance_score":80,"reason":"为什么匹配"}}]}}"""
        try:
            response = await client.chat.completions.create(
                model=RESEARCH_PRO_MODEL,
                max_tokens=3072,
                temperature=0.1,
                messages=[
                    {"role": "system", "content": "你是案例检索筛选助手，只输出可解析 JSON。"},
                    {"role": "user", "content": prompt},
                ],
            )
            selected = parse_json_object(response.choices[0].message.content or "").get("selected") or []
        except Exception:
            selected = []

        cleaned = []
        for pick_index, pick in enumerate(selected, 1):
            try:
                score = int(pick.get("relevance_score", 0))
            except Exception:
                score = 0
            score = max(0, min(100, score))
            if score < min_keep_score:
                continue
            picked_scid = str(pick.get("scid") or "未提取到信息").strip() or "未提取到信息"
            picked_ordinary_id = str(pick.get("ordinary_id") or (pick.get("id") if search_channel == "ordinary" else "") or "未提取到信息").strip() or "未提取到信息"
            picked_authority_id = str(pick.get("authority_id") or (pick.get("id") if search_channel == "authoritative" else "") or "未提取到信息").strip() or "未提取到信息"
            matched_by = [str(value) for value in (pick.get("matched_by") or []) if str(value) in ("semantic", "ordinary", "authoritative")]
            if search_channel not in matched_by:
                matched_by.append(search_channel)
            cleaned.append({
                "source_id": f"case-{idx}-{pick_index}",
                "search_index": idx,
                "type": "case",
                "scid": picked_scid,
                "ordinary_id": picked_ordinary_id,
                "authority_id": picked_authority_id,
                "title": str(pick.get("title") or "未提取到信息").strip() or "未提取到信息",
                "case_no": str(pick.get("case_no") or "未提取到信息").strip() or "未提取到信息",
                "court": str(pick.get("court") or "未提取到信息").strip() or "未提取到信息",
                "link": str(pick.get("link") or "未提取到信息").strip() or "未提取到信息",
                "case_library": search_channel,
                "authority_type": str(pick.get("authority_type") or "未提取到信息").strip() or "未提取到信息",
                "matched_by": matched_by,
                "query": item.get("query") or "",
                "reason": str(pick.get("reason") or "模型筛选匹配案例")[:300],
                "relevance_score": score,
                "raw_result": raw_result[:14000],
            })
            if len(cleaned) >= strength["limit"]:
                break
        return cleaned

    def dedupe_selected_scids(selected: list[dict]) -> list[dict]:
        def clean_case_no(value: str) -> str:
            return str(value or "").replace("(", "（").replace(")", "）").replace(" ", "").strip()

        merged: dict[str, dict] = {}
        for item in sorted(selected, key=lambda x: x.get("relevance_score", 0), reverse=True):
            scid = str(item.get("scid") or "").strip()
            case_no = clean_case_no(item.get("case_no"))
            title = str(item.get("title") or "").strip()
            ordinary_id = str(item.get("ordinary_id") or "").strip()
            authority_id = str(item.get("authority_id") or "").strip()
            key = case_no if case_no and case_no != "未提取到信息" else next(
                (value for value in (ordinary_id, authority_id, scid, title) if value and value != "未提取到信息"),
                "",
            )
            if not key:
                key = item.get("source_id")
            if key not in merged:
                item["case_no"] = case_no or item.get("case_no")
                merged[key] = item
                continue
            existing = merged[key]
            for field in ("scid", "ordinary_id", "authority_id", "authority_type", "link", "court"):
                if existing.get(field) in (None, "", "未提取到信息") and item.get(field) not in (None, "", "未提取到信息"):
                    existing[field] = item[field]
            existing["matched_by"] = list(dict.fromkeys((existing.get("matched_by") or []) + (item.get("matched_by") or [])))
            existing["relevance_score"] = max(existing.get("relevance_score", 0), item.get("relevance_score", 0))

        deduped = list(merged.values())
        for item in deduped:
            channels = set(item.get("matched_by") or [])
            if "ordinary" in channels and "authoritative" in channels:
                item["case_library"] = "mixed"
            elif "authoritative" in channels:
                item["case_library"] = "authoritative"
            elif "ordinary" in channels:
                item["case_library"] = "ordinary"
            else:
                item["case_library"] = "semantic"
        return sorted(deduped, key=lambda x: x.get("relevance_score", 0), reverse=True)[:strength["limit"]]

    async def select_candidate_laws(client, item: dict, idx: int) -> list[dict]:
        raw_result = str(item.get("result") or "")
        if not result_has_material(raw_result):
            return []
        prompt = f"""请从 MCP 法规/法条检索结果中筛选与用户问题匹配的法条，只输出 JSON。
用户问题：{query_text}
地域限制：{region or "无"}
检索词：{item.get("query") or ""}
检索目的：{item.get("reason") or ""}
筛选强度：{strength["name"]}，相关度阈值约 {strength["threshold"]}。{strength["instruction"]}

评分口径：
- 90-100：条文直接覆盖用户问题的请求权基础、责任构成或抗辩条件。
- 80-89：条文虽非唯一依据，但能直接支撑核心法律分析。
- 70-79：条文对应的制度、规则或解释路径可用于回答用户问题。
- 60-69：背景规范相关，需要结合其他依据谨慎适用。
- 50-59：弱相关，仅作补充线索。
- 50 以下：通常不展示，除非宽松模式下有检索线索价值。
不要因为条文标题、效力信息、链接或部分字段缺失而机械降分；只要该法条的制度规则可用于回答用户问题，应给 70 分以上。

MCP 原始返回：
{raw_result[:14000]}

任务：
1. 只筛选与用户问题相关的法规/法条。
2. 必须尽量提取法规名称、条号、条文内容、链接或详情链接；缺失写“未提取到信息”。
3. 如果一个结果包含多个相关条款，可以拆成多条输出。
4. 给出 relevance_score 0-100 和筛选理由。
5. relevance_score 低于 {strength["threshold"]} 的通常不要输出，除非具有明显参考价值。
6. 最多输出 {strength["limit"]} 条。

JSON：
{{"selected":[{{"law_name":"法规名称或未提取到信息","article_no":"条号或未提取到信息","title":"展示标题","article_text":"条文内容或未提取到信息","effective_status":"效力/时效或未提取到信息","link":"链接或未提取到信息","relevance_score":80,"reason":"为什么匹配"}}]}}"""
        try:
            response = await client.chat.completions.create(
                model=RESEARCH_PRO_MODEL,
                max_tokens=3072,
                temperature=0.1,
                messages=[
                    {"role": "system", "content": "你是法条检索筛选助手，只输出可解析 JSON。"},
                    {"role": "user", "content": prompt},
                ],
            )
            selected = parse_json_object(response.choices[0].message.content or "").get("selected") or []
        except Exception:
            selected = []

        cleaned = []
        for pick_index, pick in enumerate(selected, 1):
            try:
                score = int(pick.get("relevance_score", 0))
            except Exception:
                score = 0
            score = max(0, min(100, score))
            if score < min_keep_score:
                continue
            law_name = str(pick.get("law_name") or pick.get("法规名称") or "未提取到信息").strip() or "未提取到信息"
            article_no = str(pick.get("article_no") or pick.get("条号") or "未提取到信息").strip() or "未提取到信息"
            title = str(pick.get("title") or f"{law_name} {article_no}").strip() or "法条材料"
            cleaned.append({
                "source_id": f"law-{idx}-{pick_index}",
                "search_index": idx,
                "type": "law",
                "law_name": law_name,
                "article_no": article_no,
                "title": title[:140],
                "article_text": str(pick.get("article_text") or pick.get("条文内容") or "未提取到信息").strip() or "未提取到信息",
                "effective_status": str(pick.get("effective_status") or pick.get("效力/时效") or "未提取到信息").strip() or "未提取到信息",
                "link": str(pick.get("link") or pick.get("链接") or "未提取到信息").strip() or "未提取到信息",
                "query": item.get("query") or "",
                "reason": str(pick.get("reason") or "模型筛选匹配法条")[:300],
                "relevance_score": score,
                "raw_result": raw_result[:14000],
            })
            if len(cleaned) >= strength["limit"]:
                break
        return cleaned

    def dedupe_selected_laws(selected: list[dict]) -> list[dict]:
        seen = set()
        deduped = []
        for item in sorted(selected, key=lambda x: x.get("relevance_score", 0), reverse=True):
            law_name = str(item.get("law_name") or "").strip()
            article_no = str(item.get("article_no") or "").strip()
            link = str(item.get("link") or "").strip()
            title = str(item.get("title") or "").strip()
            key_parts = [
                law_name if law_name and law_name != "未提取到信息" else "",
                article_no if article_no and article_no != "未提取到信息" else "",
                link if link and link != "未提取到信息" else "",
            ]
            key = "|".join(key_parts).strip("|") or title or item.get("source_id")
            if key in seen:
                continue
            seen.add(key)
            deduped.append(item)
            if len(deduped) >= strength["limit"]:
                break
        return deduped

    async def structure_selected_cases(client, selected_cases: list[dict]) -> list[dict]:
        if not selected_cases:
            return []
        detail_semaphore = asyncio.Semaphore(4)

        def decode_json(value):
            if isinstance(value, str):
                try:
                    return decode_json(json.loads(value))
                except (TypeError, ValueError, json.JSONDecodeError):
                    return value
            if isinstance(value, list):
                return [decode_json(item) for item in value]
            if isinstance(value, dict):
                return {key: decode_json(item) for key, item in value.items()}
            return value

        def find_detail_rows(value) -> list[dict]:
            value = decode_json(value)
            if isinstance(value, list):
                if value and isinstance(value[0], dict) and any(key in value[0] for key in ("id", "ah", "title", "content")):
                    return value
                rows = []
                for item in value:
                    rows.extend(find_detail_rows(item))
                return rows
            if isinstance(value, dict):
                if isinstance(value.get("dataPreview"), (dict, list, str)):
                    rows = find_detail_rows(value["dataPreview"])
                    if rows:
                        return rows
                if "data" in value:
                    rows = find_detail_rows(value["data"])
                    if rows:
                        return rows
                for nested in value.values():
                    rows = find_detail_rows(nested)
                    if rows:
                        return rows
            return []

        def first_value(*values, default="未提取到信息"):
            for value in values:
                if isinstance(value, list):
                    value = "、".join(str(item) for item in value if item not in (None, ""))
                text = str(value or "").strip()
                if text and text not in ("None", "null", "未提取到信息"):
                    return text
            return default

        async def structure_one(idx: int, item: dict) -> dict | None:
            async def fetch_detail(*, case_id: str = "", case_no: str = "", case_type: str = "") -> str:
                async with detail_semaphore:
                    return await get_case_details_direct(case_id=case_id, case_no=case_no, case_type=case_type)

            detail_calls = []
            detail_labels = []
            if item.get("ordinary_id") not in (None, "", "未提取到信息"):
                detail_calls.append(fetch_detail(case_id=item["ordinary_id"], case_type="ptal"))
                detail_labels.append("普通案例详情")
            if item.get("authority_id") not in (None, "", "未提取到信息"):
                detail_calls.append(fetch_detail(case_id=item["authority_id"], case_type="qwal"))
                detail_labels.append("权威案例详情")
            if not detail_calls and item.get("case_no") not in (None, "", "未提取到信息"):
                detail_calls.append(fetch_detail(case_no=item["case_no"]))
                detail_labels.append("按案号查询详情")
            elif not detail_calls and item.get("scid") not in (None, "", "未提取到信息"):
                detail_calls.append(fetch_detail(case_id=item["scid"]))
                detail_labels.append("语义检索案例详情")

            detail_results = await asyncio.gather(*detail_calls, return_exceptions=True) if detail_calls else []
            records = []
            for result in detail_results:
                if not isinstance(result, Exception):
                    records.extend(find_detail_rows(result))

            primary = records[0] if records else {}
            field_labels = {
                "id": "案例 ID", "type": "案例类型", "ah": "案号", "title": "案名",
                "jbdw": "法院", "ajlb": "案件类别", "ajlx": "案件类型", "spcx": "审判程序",
                "wszl": "文书种类", "ay": "案由", "cprq": "裁判日期", "xzqh_p": "省份",
                "xzqh_c": "城市", "yyft": "援引法条", "zyjd": "争议焦点", "ayjd": "争议焦点",
                "ajjbqk": "案件基本情况", "ajjbqk_zj": "证据情况", "ajjbqk_bh": "辩护情况",
                "ajjbqk_zk": "指控情况", "ajjbqk_bssl": "本审审理情况", "dsr": "当事人",
                "ssjl": "诉讼记录", "cmss": "查明事实", "fxgc": "分析过程", "pjjg": "裁判结果",
                "section": "结构化段落", "judge": "法官", "lawyer": "律师", "url": "链接",
            }
            detail_fields = {
                field_labels.get(key, key): value
                for key, value in primary.items()
                if key != "content" and value not in (None, "", [], {})
            }
            link = first_value(primary.get("url"), item.get("link"))
            detail_fields.setdefault("scid", item.get("scid") or "未提取到信息")
            detail_fields.setdefault("检索来源", "、".join(item.get("matched_by") or []) or "semantic")
            ref = {
                "id": f"detail-case-{idx}",
                "source_id": item.get("source_id"),
                "type": "case",
                "title": first_value(primary.get("title"), item.get("title")),
                "query": item.get("query") or "",
                "link": link,
                "relevance_score": item.get("relevance_score") or 0,
                "relevance_reason": item.get("reason") or "检索结果与用户问题相关",
                "fields": detail_fields,
                "content": first_value(primary.get("content"), item.get("raw_result")),
                "case_library": item.get("case_library") or "semantic",
                "authority_type": item.get("authority_type") or "",
                "matched_by": item.get("matched_by") or [],
                "scid": item.get("scid") or "",
                "ordinary_id": item.get("ordinary_id") or "",
                "authority_id": item.get("authority_id") or "",
                "curated": True,
            }
            return ref

        structured = await asyncio.gather(
            *(structure_one(idx, item) for idx, item in enumerate(selected_cases, 1))
        )
        structured = [item for item in structured if item]
        return sorted(structured, key=lambda r: r.get("relevance_score", 0), reverse=True)

    async def structure_selected_laws(client, selected_laws: list[dict]) -> list[dict]:
        if not selected_laws:
            return []
        semaphore = asyncio.Semaphore(4)

        async def structure_one(idx: int, item: dict) -> dict | None:
            prompt = f"""请根据已筛选法条和原文片段提取结构化法条信息，只输出 JSON。
用户问题：{query_text}
地域限制：{region or "无"}

已筛选法条：
{json.dumps(item, ensure_ascii=False)[:16000]}

要求：
1. 必须输出这一条法条，不要因为字段缺失而丢弃。
2. 缺失字段统一写“未提取到信息”。
3. 必须保留法规名称、条号、链接字段。
4. 不要编造原文没有的信息。
5. content 字段请写成可读摘要，说明该条文内容及其对用户问题的适用关系。

JSON：
{{"reference":{{"source_id":"{item.get("source_id") or ""}","type":"law","title":"法规名称 + 条号","query":"检索词","link":"链接或未提取到信息","relevance_score":80,"relevance_reason":"相关性说明","fields":{{"法规名称":"未提取到信息","条号":"未提取到信息","条文内容":"未提取到信息","效力/时效":"未提取到信息","适用关系":"未提取到信息","链接":"未提取到信息"}},"content":"完整法条摘要"}}}}"""
            ref = {}
            try:
                async with semaphore:
                    response = await client.chat.completions.create(
                        model=DEEPSEEK_MODEL,
                        max_tokens=2200,
                        temperature=0.1,
                        messages=[
                            {"role": "system", "content": "你是法条结构化提取助手，只输出可解析 JSON。"},
                            {"role": "user", "content": prompt},
                        ],
                    )
                data = parse_json_object(response.choices[0].message.content or "")
                ref = data.get("reference") or (data.get("references") or [{}])[0] or {}
            except Exception:
                ref = {}

            if not ref:
                ref = {
                    "source_id": item.get("source_id"),
                    "type": "law",
                    "title": item.get("title") or f"{item.get('law_name') or '未提取到信息'} {item.get('article_no') or ''}".strip(),
                    "query": item.get("query") or "",
                    "link": item.get("link") or "未提取到信息",
                    "relevance_score": item.get("relevance_score") or 0,
                    "relevance_reason": item.get("reason") or "模型筛选匹配法条",
                    "fields": {
                        "法规名称": item.get("law_name") or "未提取到信息",
                        "条号": item.get("article_no") or "未提取到信息",
                        "条文内容": item.get("article_text") or "未提取到信息",
                        "效力/时效": item.get("effective_status") or "未提取到信息",
                        "适用关系": item.get("reason") or "经筛选与用户问题相关，需结合案件事实适用。",
                        "链接": item.get("link") or "未提取到信息",
                    },
                    "content": (
                        f"该法条在筛选阶段被判定为相关。"
                        f"法规名称：{item.get('law_name') or '未提取到信息'}；"
                        f"条号：{item.get('article_no') or '未提取到信息'}；"
                        f"条文内容：{item.get('article_text') or '未提取到信息'}。"
                    ),
                }

            ref["type"] = "law"
            fields = ref.get("fields") if isinstance(ref.get("fields"), dict) else {}
            fields.setdefault("法规名称", item.get("law_name") or "未提取到信息")
            fields.setdefault("条号", item.get("article_no") or "未提取到信息")
            fields.setdefault("条文内容", item.get("article_text") or "未提取到信息")
            fields.setdefault("效力/时效", item.get("effective_status") or "未提取到信息")
            fields.setdefault("适用关系", item.get("reason") or "未提取到信息")
            fields.setdefault("链接", item.get("link") or "未提取到信息")
            ref["fields"] = fields
            ref["source_id"] = ref.get("source_id") or item.get("source_id")
            ref["query"] = ref.get("query") or item.get("query") or ""
            ref["link"] = ref.get("link") or item.get("link") or "未提取到信息"
            ref["relevance_score"] = ref.get("relevance_score") or item.get("relevance_score") or 0
            ref["relevance_reason"] = ref.get("relevance_reason") or item.get("reason") or ""
            return normalize_reference(ref, idx, prefix="law-final")

        structured = await asyncio.gather(
            *(structure_one(idx, item) for idx, item in enumerate(selected_laws, 1))
        )
        structured = [item for item in structured if item]
        return sorted(structured, key=lambda r: r.get("relevance_score", 0), reverse=True)

    async def dedupe_and_rank_references(client, candidates: list[dict]) -> list[dict]:
        if not candidates:
            return []
        source_text = json.dumps(candidates[:40], ensure_ascii=False)[:30000]
        prompt = f"""请对候选法律材料做全局去重、合并和相关性排序，只输出 JSON。
用户问题：{query_text}
地域限制：{region or "无"}
输出侧重：{focus}

候选材料：
{source_text}

要求：
1. 对重复案例/法条去重；同案号、同案名、同一链接、内容高度相似均视为可合并。
2. 判断与用户问题的相关性，relevance_score 为 0-100。
3. 优先保留地域、案由、事实、裁判观点更贴近用户问题的材料。
4. 缺失字段统一写“未提取到信息”，不要编造。
5. 必须保留或合并链接字段；没有链接写“未提取到信息”。
6. 每条输出必须保留最主要来源候选的 source_id。
7. 最多输出 {strength["limit"]} 条。

JSON：
{{"references":[{{"source_id":"候选材料 source_id","type":"case","title":"案例标题","query":"检索词","link":"链接或未提取到信息","relevance_score":80,"relevance_reason":"相关性说明","fields":{{"案名":"未提取到信息","法院":"未提取到信息","案号":"未提取到信息","案由":"未提取到信息","裁判日期":"未提取到信息","核心事实":"...","争议焦点":"...","裁判观点":"...","参考价值":"...","链接":"未提取到信息"}},"content":"去重整理后的完整材料"}}]}}"""
        try:
            response = await client.chat.completions.create(
                model=DEEPSEEK_MODEL,
                max_tokens=4096,
                temperature=0.1,
                messages=[
                    {"role": "system", "content": "你是严谨的法律检索去重排序助手，只输出可解析 JSON。"},
                    {"role": "user", "content": prompt},
                ],
            )
            refs = parse_json_object(response.choices[0].message.content or "").get("references") or []
        except Exception:
            refs = []

        cleaned = []
        for idx, ref in enumerate(refs, 1):
            normalized = normalize_reference(ref, idx, prefix="final")
            if normalized:
                cleaned.append(normalized)
        repaired = repair_final_references(cleaned, candidates)
        return sorted(repaired, key=lambda r: r.get("relevance_score", 0), reverse=True)

    def is_missing_value(value) -> bool:
        text = str(value or "").strip()
        return not text or text in ("未提取到信息", "检索结果未显示", "None", "null")

    def repair_final_references(final_refs: list[dict], candidates: list[dict]) -> list[dict]:
        candidate_map = {c.get("source_id") or c.get("id"): c for c in candidates}
        repaired = []
        for ref in final_refs:
            source = candidate_map.get(ref.get("source_id"))
            if not source:
                # 次优匹配：同类型同 query 的候选
                source = next(
                    (
                        c for c in candidates
                        if c.get("type") == ref.get("type")
                        and c.get("query")
                        and c.get("query") == ref.get("query")
                    ),
                    None,
                )
            if source:
                ref_fields = ref.get("fields") or {}
                source_fields = source.get("fields") or {}
                for key, value in source_fields.items():
                    if is_missing_value(ref_fields.get(key)) and not is_missing_value(value):
                        ref_fields[key] = value
                ref["fields"] = ref_fields
                if is_missing_value(ref.get("title")) or ref.get("title") in ("案例材料", "法条材料"):
                    ref["title"] = source.get("title") or ref.get("title")
                if is_missing_value(ref.get("content")):
                    ref["content"] = source.get("content") or ref.get("content")
                if is_missing_value(ref.get("link")) and not is_missing_value(source.get("link")):
                    ref["link"] = source.get("link")
                    ref["fields"]["链接"] = source.get("link")
            repaired.append(ref)
        return repaired

    def result_has_material(text: str) -> bool:
        stripped = (text or "").strip()
        if not stripped:
            return False
        lowered = stripped.lower()
        empty_markers = {
            "未找到",
            "无结果",
            "没有结果",
            "未找到案例检索工具",
            "未找到法律法规检索工具",
            "工具执行完成，无文本输出",
        }
        if any(marker in lowered for marker in empty_markers):
            return False
        try:
            data = json.loads(stripped)
            if isinstance(data, dict):
                err = data.get("error")
                if err and str(err).strip() not in ("0", "None", "null", "false", "False"):
                    return False
                # 问题2修复: 递归搜索嵌套数据容器
                return _has_nested_material(data)
            if isinstance(data, list):
                return len(data) > 0
        except Exception:
            pass
        return True

    def _has_nested_material(obj, depth: int = 0) -> bool:
        """递归检查 JSON 对象中是否包含实质数据内容（问题2修复）"""
        if depth > 8:
            return False
        if isinstance(obj, dict):
            # 检查 _extract_nested_content 产生的富化结构
            items = obj.get("items")
            if isinstance(items, list) and len(items) > 0:
                return True
            # 检查标准数据容器 key
            for key in ("data", "result", "results", "items", "list", "content", "text",
                        "wenshu", "cases", "laws", "records", "extra", "dataPreview"):
                value = obj.get(key)
                if value:
                    if isinstance(value, (dict, list)):
                        if _has_nested_material(value, depth + 1):
                            return True
                    elif isinstance(value, str) and len(value.strip()) > 10:
                        return True
                    elif value:
                        return True
            return bool(obj)
        if isinstance(obj, list):
            if len(obj) > 0 and isinstance(obj[0], dict):
                return True
            return len(obj) > 0
        return bool(obj)

    def raw_result_reference(item: dict, idx: int, prefix: str = "raw") -> dict | None:
        raw = str(item.get("result") or "").strip()
        if not result_has_material(raw):
            return None
        item_type = item.get("type")
        if item_type not in ("case", "law"):
            return None

        display_title = f"{'案例' if item_type == 'case' else '法条'}检索结果：{item.get('query', '')[:80]}"
        fields = {
            "检索词": item.get("query") or "未提取到信息",
            "材料来源": "元典检索原始返回",
            "链接": "未提取到信息",
        }

        parsed_title = ""
        parsed_fields = {}
        try:
            parsed = json.loads(raw)
            candidates = []

            def collect(obj):
                if isinstance(obj, dict):
                    if any(k in obj for k in ("案号", "案由", "法院", "裁判日期", "案件名称", "案名", "title", "name", "lawName", "article", "content", "url", "link")):
                        candidates.append(obj)
                    for value in obj.values():
                        collect(value)
                elif isinstance(obj, list):
                    for value in obj:
                        collect(value)

            collect(parsed)
            first = candidates[0] if candidates else (parsed if isinstance(parsed, dict) else {})
            if isinstance(first, dict):
                link = first.get("链接") or first.get("url") or first.get("link") or first.get("href") or "未提取到信息"
                if item_type == "case":
                    parsed_title = str(
                        first.get("案件名称") or first.get("案名") or first.get("title") or first.get("name") or ""
                    )
                    parsed_fields = {
                        "案名": first.get("案件名称") or first.get("案名") or first.get("title") or first.get("name") or "未提取到信息",
                        "法院": first.get("法院") or first.get("court") or "未提取到信息",
                        "案号": first.get("案号") or first.get("caseNo") or first.get("case_no") or "未提取到信息",
                        "案由": first.get("案由") or first.get("cause") or "未提取到信息",
                        "裁判日期": first.get("裁判日期") or first.get("date") or first.get("judgementDate") or "未提取到信息",
                        "链接": link,
                    }
                else:
                    parsed_title = str(
                        first.get("法规名称") or first.get("lawName") or first.get("title") or first.get("name") or ""
                    )
                    parsed_fields = {
                        "法规名称": first.get("法规名称") or first.get("lawName") or first.get("title") or first.get("name") or "未提取到信息",
                        "条号": first.get("条号") or first.get("article") or first.get("articleNo") or "未提取到信息",
                        "链接": link,
                    }
        except Exception:
            parsed_fields = {}

        if parsed_title:
            display_title = parsed_title[:120]
        fields.update({k: str(v) for k, v in parsed_fields.items() if str(v or "").strip()})
        return {
            "id": f"{prefix}-{item_type}-{idx}",
            "source_id": f"{prefix}-{item_type}-{idx}",
            "type": item_type,
            "title": display_title,
            "query": item.get("query") or "",
            "reason": "原始检索结果，字段请以全文为准。",
            "relevance_score": 30,
            "link": fields.get("链接") or "未提取到信息",
            "fields": fields,
            "content": raw[:12000],
            "raw_fallback": True,
        }

    def fallback_raw_references(results: list[dict]) -> list[dict]:
        refs = []
        for idx, item in enumerate(results, 1):
            ref = raw_result_reference(item, idx, prefix="raw")
            if ref:
                refs.append(ref)
        return refs

    def build_raw_result_context(results: list[dict]) -> str:
        parts = []
        for idx, item in enumerate(results, 1):
            raw = str(item.get("result") or "").strip()
            if not result_has_material(raw):
                continue
            parts.append(
                f"[原始检索 {idx}] type={item.get('type')}\n"
                f"query={item.get('query')}\n"
                f"reason={item.get('reason')}\n"
                f"result={raw[:8000]}"
            )
        return "\n\n".join(parts) or "无可用原始检索结果。"

    def build_reference_context(refs: list[dict]) -> str:
        if not refs:
            return "未形成结构化材料。请结合原始检索结果判断，不要直接认定检索不到。"
        parts = []
        for idx, ref in enumerate(refs, 1):
            fields = ref.get("fields") or {}
            field_text = "\n".join(f"{k}: {v}" for k, v in fields.items())
            parts.append(
                f"[{idx}] {ref.get('type')} {ref.get('title')}\n"
                f"相关度：{ref.get('relevance_score')}\n"
                f"相关性：{ref.get('reason')}\n"
                f"{field_text}\n{ref.get('content') or ''}"
            )
        return "\n\n".join(parts)

    async def generate():
        answer_text = ""
        references = []
        planned_results = []
        candidates = []
        selected_cases = []
        selected_laws = []
        try:
            if not is_mcp_available():
                yield sse_payload({"status": "元典 MCP 未启用，无法检索具体法规/案例。", "tool": {"stage": "failed", "name": "检索服务"}})

            client = get_ai_client()
            yield sse_payload({"status": "正在生成检索计划...", "tool": {"stage": "start", "name": "检索计划"}})
            plan = await create_plan(client)
            channel_plan_counts = {
                channel: sum(1 for item in plan if item.get("type") == "case" and item.get("channel") == channel)
                for channel in ("semantic", "ordinary", "authoritative")
            }
            plan_summary = "、".join(f"{channel} {count} 个" for channel, count in channel_plan_counts.items() if count)
            yield sse_payload({
                "status": f"已生成 {len(plan)} 个检索任务" + (f"（{plan_summary}）" if plan_summary else "") + "。",
                "tool": {"stage": "done", "name": "检索计划"},
            })

            mcp_semaphore = asyncio.Semaphore(4)
            llm_semaphore = asyncio.Semaphore(4)

            async def run_research_task(idx: int, item: dict) -> dict:
                kind = "案例检索" if item["type"] == "case" else "法规检索"
                events = [{
                    "status": f"正在执行第 {idx} 个任务：{kind} - {item['query']}",
                    "tool": {"stage": "start", "name": kind, "query": item["query"]},
                }]
                async def guarded(call):
                    async with mcp_semaphore:
                        return await call

                if item["type"] == "case":
                    channel = item.get("channel") or "semantic"
                    if channel == "ordinary":
                        raw_result = await guarded(search_case_ordinary_direct(dict(item.get("arguments") or {}), region=region, date_start=date_start, date_end=date_end))
                        result = "[普通案例 ptal_search]\n" + str(raw_result or "")[:12000]
                    elif channel == "authoritative":
                        raw_result = await guarded(search_case_authoritative_direct(dict(item.get("arguments") or {}), region=region, date_start=date_start, date_end=date_end))
                        result = "[权威案例 qwal_search]\n" + str(raw_result or "")[:12000]
                    else:
                        raw_result = await guarded(search_case_vector_direct(item["query"], region=region, return_num=15, date_start=date_start, date_end=date_end))
                        result = "[语义检索 vector_search]\n" + str(raw_result or "")[:12000]
                    events.append({
                        "status": f"第 {idx} 个案例任务完成：{channel} 通道。",
                        "tool": {"stage": "done", "name": f"案例检索-{channel}", "query": item["query"]},
                    })
                    planned = {**item, "result": result}
                else:
                    result = await guarded(search_law_direct(item["query"], region=region))
                    planned = {**item, "result": str(result or "")}
                events.append({
                    "status": f"第 {idx} 个任务完成。",
                    "tool": {"stage": "done", "name": kind, "query": item["query"]},
                })
                if item["type"] == "case":
                    events.append({
                        "status": f"正在从第 {idx} 个案例结果中筛选匹配 scid...",
                        "tool": {"stage": "start", "name": "案例筛选", "query": item["query"]},
                    })
                    async with llm_semaphore:
                        selected = await select_candidate_scids(client, planned, idx)
                    events.append({
                        "status": f"第 {idx} 个结果筛选到 {len(selected)} 条匹配案例。",
                        "tool": {"stage": "done", "name": "案例筛选", "query": item["query"]},
                    })
                    return {"planned": planned, "selected": selected, "selected_laws": [], "candidates": [], "events": events}

                events.append({
                    "status": f"正在从第 {idx} 个法条结果中筛选匹配条款...",
                    "tool": {"stage": "start", "name": "法条筛选", "query": item["query"]},
                })
                async with llm_semaphore:
                    selected_law_items = await select_candidate_laws(client, planned, idx)
                events.append({
                    "status": f"第 {idx} 个结果筛选到 {len(selected_law_items)} 条匹配法条。",
                    "tool": {"stage": "done", "name": "法条筛选", "query": item["query"]},
                })
                return {"planned": planned, "selected": [], "selected_laws": selected_law_items, "candidates": [], "events": events}

            tasks = [
                asyncio.create_task(run_research_task(idx, item))
                for idx, item in enumerate(plan, 1)
            ]
            for task in asyncio.as_completed(tasks):
                outcome = await task
                planned_results.append(outcome["planned"])
                selected_cases.extend(outcome["selected"])
                selected_laws.extend(outcome.get("selected_laws", []))
                candidates.extend(outcome["candidates"])
                for event in outcome["events"]:
                    yield sse_payload(event)

            deduped_cases = dedupe_selected_scids(selected_cases)
            yield sse_payload({"status": f"正在按 scid 去重并结构化 {len(deduped_cases)} 条案例...", "tool": {"stage": "start", "name": "案例结构化"}})
            case_refs = await structure_selected_cases(client, deduped_cases)
            yield sse_payload({"status": f"案例结构化完成，形成 {len(case_refs)} 条案例材料。", "tool": {"stage": "done", "name": "案例结构化"}})

            deduped_laws = dedupe_selected_laws(selected_laws)
            yield sse_payload({"status": f"正在按法规名、条号和链接去重并结构化 {len(deduped_laws)} 条法条...", "tool": {"stage": "start", "name": "法条结构化"}})
            law_refs = await structure_selected_laws(client, deduped_laws)
            yield sse_payload({"status": f"法条结构化完成，形成 {len(law_refs)} 条法条材料。", "tool": {"stage": "done", "name": "法条结构化"}})
            case_fallback_refs = []
            law_fallback_refs = []
            if need_cases and not case_refs:
                case_fallback_refs = fallback_raw_references([item for item in planned_results if item.get("type") == "case"])
                if case_fallback_refs:
                    yield sse_payload({"status": f"案例筛选未形成结构化卡片，已保留 {len(case_fallback_refs)} 条原始案例结果供复核。", "tool": {"stage": "done", "name": "案例兜底"}})
            if need_laws and not law_refs:
                law_fallback_refs = fallback_raw_references([item for item in planned_results if item.get("type") == "law"])
                if law_fallback_refs:
                    yield sse_payload({"status": f"法条筛选未形成结构化卡片，已保留 {len(law_fallback_refs)} 条原始法条结果供复核。", "tool": {"stage": "done", "name": "法条兜底"}})
            references = sorted(case_refs + law_refs + case_fallback_refs + law_fallback_refs, key=lambda r: r.get("relevance_score", 0), reverse=True)
            if not references:
                references = fallback_raw_references(planned_results)
            yield sse_payload({
                "references": references,
                "replace_references": True,
                "status": f"材料整理完成，形成 {len(references)} 条材料卡片。",
                "tool": {"stage": "done", "name": "材料整理"},
            })

            answer_prompt = f"""请基于检索材料回答用户问题。

用户问题：
{query_text}

地域限制：{region or "无"}
输出侧重：{focus}

整理后的检索材料：
{build_reference_context(references)}

原始检索结果：
{build_raw_result_context(planned_results)}

要求：
1. 先给结论，再说明理由。
2. 如果结构化材料为空但原始检索结果有内容，应说明“结构化抽取不足”，并基于原始检索结果谨慎分析；不要直接说检索不到。
3. 案例和法条字段缺失时，不要补编。
4. 需要回答用户提出的实体问题，而不是只罗列材料。"""
            stream_kwargs = {
                "model": RESEARCH_PRO_MODEL,
                "max_tokens": 4096,
                "temperature": 0.35,
                "stream": True,
                "stream_options": {"include_usage": True},
                "messages": [
                    {"role": "system", "content": "你是严谨的法律研究助手，请用中文输出。"},
                    {"role": "user", "content": answer_prompt},
                ],
            }
            try:
                stream = await client.chat.completions.create(**stream_kwargs)
            except Exception as stream_error:
                if "stream_options" not in str(stream_error):
                    raise
                stream_kwargs.pop("stream_options", None)
                stream = await client.chat.completions.create(**stream_kwargs)
            answer_usage = None
            async for chunk in stream:
                if getattr(chunk, "usage", None):
                    answer_usage = extract_usage(chunk.usage)
                if not chunk.choices:
                    continue
                delta = chunk.choices[0].delta
                if delta.content:
                    answer_text += delta.content
                    yield sse_payload({"content": delta.content})
            await record_usage(user=user, feature="research", model=RESEARCH_PRO_MODEL, usage=answer_usage)

            record_id = await save_research_record(
                user_id=user["id"],
                query_text=display_query_text,
                answer_text=answer_text,
                references=references,
                meta={
                    "region": region,
                    "focus": focus,
                    "filter_strength": filter_strength,
                    "need_cases": need_cases,
                    "need_laws": need_laws,
                    "plan": plan,
                    "case_channel_task_counts": channel_plan_counts,
                    "case_channel_material_counts": {
                        channel: sum(
                            1 for item in planned_results
                            if item.get("type") == "case" and item.get("channel") == channel
                            and result_has_material(str(item.get("result") or ""))
                        )
                        for channel in ("semantic", "ordinary", "authoritative")
                    },
                    "candidate_count": len(candidates),
                    "selected_case_count": len(selected_cases),
                    "selected_law_count": len(selected_laws),
                },
                messages=[
                    {"role": "user", "content": display_query_text, "files": summarize_uploaded_files(files)},
                    {"role": "assistant", "content": answer_text, "files": []},
                ],
            )
            yield sse_payload({"record_id": record_id})
            yield "data: [DONE]\n\n"
        except Exception as e:
            yield sse_payload({"error": str(e)})

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@app.post("/api/document/generate")
async def document_generate(request: Request):
    from db import get_session
    from ai_client import DEEPSEEK_MODEL, generate_document
    user = await require_current_user(request)
    token = request.cookies.get("session_token")
    session = await get_session(token, user["id"]) if token else None
    if not session:
        raise HTTPException(status_code=401, detail="请先登录")

    try:
        body = await request.json()
        doc_type = body.get("doc_type", "")
        info = body.get("info", {})
        files = body.get("files", []) or []
    except Exception:
        raise HTTPException(status_code=400, detail="无效的请求")

    if not doc_type:
        raise HTTPException(status_code=400, detail="请选择文书类型")

    try:
        if files:
            material_parts = []
            for index, file_info in enumerate(files[:8], 1):
                try:
                    material_parts.append(f"【资料{index}】\n{build_file_description(file_info, max_chars=12000)}")
                except Exception as file_error:
                    filename = file_info.get("name", f"资料{index}") if isinstance(file_info, dict) else f"资料{index}"
                    material_parts.append(f"【资料{index}: {filename}】\n（读取失败：{file_error}）")
            if material_parts:
                info = dict(info or {})
                info["用户上传的案件资料"] = "\n\n".join(material_parts)
                info["资料使用要求"] = (
                    "请优先依据上传资料提取当事人、时间线、金额、违约事实、证据和证明目的；"
                    "资料中无法确认的信息用【待补充】标注，不要编造。"
                )
        doc_content, usage_obj = await generate_document(doc_type, info, include_usage=True)
        await record_usage(
            user=user,
            feature="document_generate",
            model=DEEPSEEK_MODEL,
            usage=extract_usage(usage_obj),
        )
        return {"success": True, "content": doc_content}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"生成失败: {str(e)}")


@app.post("/api/document/extract-info")
async def document_extract_info(request: Request):
    from db import get_session
    from ai_client import get_ai_client, DEEPSEEK_MODEL, DOCUMENT_TYPES

    user = await require_current_user(request)
    token = request.cookies.get("session_token")
    session = await get_session(token, user["id"]) if token else None
    if not session:
        raise HTTPException(status_code=401, detail="请先登录")

    try:
        body = await request.json()
        doc_type = body.get("doc_type", "")
        fields = body.get("fields", []) or []
        files = body.get("files", []) or []
        current_info = body.get("current_info", {}) or {}
    except Exception:
        raise HTTPException(status_code=400, detail="无效的请求")

    if not files:
        raise HTTPException(status_code=400, detail="请先上传案件资料")

    doc_meta = DOCUMENT_TYPES.get(doc_type, {})
    doc_name = doc_meta.get("name", doc_type or "法律文书")
    allowed_fields = [str(field).strip() for field in fields if str(field).strip()]
    if not allowed_fields:
        allowed_fields = doc_meta.get("fields", [])

    try:
        material_parts = []
        sources = []
        for index, file_info in enumerate(files[:8], 1):
            filename = file_info.get("name", f"资料{index}") if isinstance(file_info, dict) else f"资料{index}"
            sources.append(filename)
            try:
                material_parts.append(f"【资料{index}: {filename}】\n{build_file_description(file_info, max_chars=14000)}")
            except Exception as file_error:
                material_parts.append(f"【资料{index}: {filename}】\n（读取失败：{file_error}）")

        materials_text = "\n\n".join(material_parts)
        prompt = f"""请从用户上传的案件资料中，为“{doc_name}”提取结构化起草信息。

只能提取资料中可以确认或高度推断的信息；无法确认的字段留空字符串，不要编造。
如果资料中存在矛盾、缺页、图片无法识别、金额/日期不清楚，请写入 notes。

需要提取的字段：
{json.dumps(allowed_fields, ensure_ascii=False)}

用户已经填写的信息（如与资料一致可参考，如冲突请在 notes 说明）：
{json.dumps(current_info, ensure_ascii=False)}

案件资料：
{materials_text}

请只输出 JSON，不要输出 Markdown。格式如下：
{{
  "info": {{
    "字段名": "提取到的内容"
  }},
  "notes": ["需要用户核对或补充的事项"]
}}"""

        client = get_ai_client()
        response = await client.chat.completions.create(
            model=DEEPSEEK_MODEL,
            max_tokens=4096,
            temperature=0.2,
            messages=[
                {"role": "system", "content": "你是法律文书资料整理助手，只输出可解析 JSON。"},
                {"role": "user", "content": prompt},
            ],
        )
        await record_response_usage(user=user, feature="document_extract", model=DEEPSEEK_MODEL, response=response)
        raw = response.choices[0].message.content or "{}"
        raw = raw.strip()
        if raw.startswith("```"):
            raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
            raw = re.sub(r"\s*```$", "", raw)
        try:
            data = json.loads(raw)
        except json.JSONDecodeError:
            match = re.search(r"\{[\s\S]*\}", raw)
            data = json.loads(match.group(0)) if match else {}

        extracted = data.get("info", {}) if isinstance(data, dict) else {}
        if not isinstance(extracted, dict):
            extracted = {}
        cleaned = {}
        for field in allowed_fields:
            value = extracted.get(field, "")
            if isinstance(value, (list, tuple)):
                value = "\n".join(str(item).strip() for item in value if str(item).strip())
            cleaned[field] = str(value or "").strip()

        notes = data.get("notes", []) if isinstance(data, dict) else []
        if isinstance(notes, str):
            notes = [notes]
        if not isinstance(notes, list):
            notes = []

        return {
            "success": True,
            "info": cleaned,
            "notes": [str(note).strip() for note in notes if str(note).strip()][:8],
            "sources": sources,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"提取失败: {str(e)}")


@app.post("/api/document/export")
async def document_export(request: Request):
    from db import get_session
    from document_templates import export_to_word
    user = await require_current_user(request)
    token = request.cookies.get("session_token")
    session = await get_session(token, user["id"]) if token else None
    if not session:
        raise HTTPException(status_code=401, detail="请先登录")

    try:
        body = await request.json()
        content = body.get("content", "")
        title = body.get("title", "法律文书")
    except Exception:
        raise HTTPException(status_code=400, detail="无效的请求")

    try:
        filepath = export_to_word(content, title)
        return FileResponse(
            filepath,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{title}.docx",
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


def _extract_contract_json(content: str) -> dict:
    text = (content or "").strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\s*```$", "", text)
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        match = re.search(r"\{[\s\S]*\}", text)
        if match:
            return json.loads(match.group(0))
        raise


async def _extract_or_repair_contract_json(client, content: str, model: str) -> dict:
    try:
        return _extract_contract_json(content)
    except json.JSONDecodeError as original_error:
        repair_prompt = f"""下面是一段合同审查模型输出，但不是合法 JSON。请修复为合法 JSON，只输出 JSON，不要输出 Markdown 或解释。

必须保留并整理为以下结构：
{{
  "summary": "string",
  "overall_risk": "high|medium|low",
  "issues": [
    {{
      "id": "R01",
      "title": "string",
      "severity": "high|medium|low",
      "category": "string",
      "page": 1,
      "quote": "string",
      "analysis": "string",
      "basis": "string",
      "suggestion": "string"
    }}
  ]
}}

修复规则：
1. 只修复 JSON 语法，不新增事实。
2. 对字符串中的换行、引号等做必要转义。
3. 如果某个字段缺失，用空字符串或合理默认值补齐。
4. issues 必须是数组。

原始解析错误：
{original_error}

待修复内容：
{(content or "")[:12000]}"""
        response = await client.chat.completions.create(
            model=model,
            max_tokens=8192,
            temperature=0,
            messages=[
                {"role": "system", "content": "你是 JSON 修复器，只输出可解析 JSON。"},
                {"role": "user", "content": repair_prompt},
            ],
        )
        return _extract_contract_json(response.choices[0].message.content or "{}")


def _normalize_contract_review(data: dict, issue_limit: int | None = 12) -> dict:
    issues = data.get("issues") if isinstance(data, dict) else []
    if not isinstance(issues, list):
        issues = []
    if issue_limit and issue_limit > 0:
        issues = issues[:issue_limit]
    normalized = []
    for index, item in enumerate(issues, 1):
        if not isinstance(item, dict):
            continue
        severity = str(item.get("severity") or "medium").lower()
        if severity not in ("high", "medium", "low"):
            severity = "medium"
        try:
            page = int(item.get("page") or 1)
        except Exception:
            page = 1
        quote = re.sub(r"\s+", " ", str(item.get("quote") or "")).strip()
        quote = re.sub(r"^(原文|合同原文|引用|quote)[:：]\s*", "", quote, flags=re.IGNORECASE)
        quote = quote.strip("“”\"'` ")
        normalized.append({
            "id": str(item.get("id") or f"R{index:02d}"),
            "title": str(item.get("title") or "风险点"),
            "severity": severity,
            "category": str(item.get("category") or "条款风险"),
            "page": max(1, page),
            "quote": quote[:300],
            "analysis": str(item.get("analysis") or "")[:1200],
            "basis": str(item.get("basis") or "")[:1200],
            "suggestion": str(item.get("suggestion") or "")[:1200],
        })
    return {
        "summary": str(data.get("summary") or "已完成合同审查。")[:1200] if isinstance(data, dict) else "已完成合同审查。",
        "overall_risk": str(data.get("overall_risk") or "medium").lower() if isinstance(data, dict) else "medium",
        "issues": normalized,
    }


CONTRACT_REVIEW_TYPES = {
    "general": {
        "name": "通用合同",
        "focus": "全面审查主体、标的、价款、履行、违约、解除、争议解决、通知送达、附件效力等基础条款。",
    },
    "sale": {
        "name": "买卖合同",
        "focus": "重点审查标的规格、质量标准、验收流程、交付风险、价款结算、所有权转移、质量异议、违约责任和退换货安排。",
    },
    "lease": {
        "name": "租赁合同",
        "focus": "重点审查租赁物状态、租期、押金、租金支付、维修责任、转租限制、提前解除、违约金、腾退交还和装修归属。",
    },
    "labor": {
        "name": "劳动/劳务合同",
        "focus": "重点审查用工性质、岗位地点、薪酬工时、社保福利、试用期、竞业限制、保密义务、解除终止、违约责任与劳动合规风险。",
    },
    "service": {
        "name": "服务合同",
        "focus": "重点审查服务范围、交付成果、验收标准、付款节点、人员安排、知识产权、数据/保密、延期责任和服务质量救济。",
    },
    "equity": {
        "name": "股权转让/投资协议",
        "focus": "重点审查股权权属、转让价款、先决条件、工商变更、陈述保证、交割安排、税费承担、回购/对赌、公司治理和违约救济。",
    },
    "loan": {
        "name": "借款/担保合同",
        "focus": "重点审查借款用途、利率费用、还款安排、提前到期、担保范围、抵押/质押登记、保证方式、违约责任和实现债权条款。",
    },
    "nda": {
        "name": "保密/合作协议",
        "focus": "重点审查保密信息范围、例外情形、使用限制、披露对象、保密期限、知识产权归属、违约赔偿和竞业/排他限制。",
    },
}


def _get_contract_review_type(contract_type: str) -> tuple[str, str]:
    item = CONTRACT_REVIEW_TYPES.get(contract_type) or CONTRACT_REVIEW_TYPES["general"]
    return item["name"], item["focus"]


def _parse_issue_limit(value, default: int = 12) -> int:
    try:
        limit = int(value)
    except Exception:
        return default
    if limit <= 0:
        return 0
    return max(1, min(limit, 100))


async def _build_contract_reference_queries(
    *,
    fname: str,
    contract_type_name: str,
    note: str,
    description: str,
) -> list[dict]:
    from ai_client import CONTRACT_REFERENCE_QUERY_MODEL, get_ai_client

    plain = re.sub(r"\s+", " ", description)
    seed = plain[:3000]
    fallback_base = f"{contract_type_name} 合同审查 {note or ''} {plain[:500]}".strip()
    fallback_queries = [
        {
            "type": "law",
            "query": f"{fallback_base} 民法典 合同编 违约责任 解除"[:300],
            "reason": "固定兜底法规检索词",
        },
        {
            "type": "case",
            "query": f"{fallback_base} 合同纠纷 裁判规则"[:300],
            "reason": "固定兜底案例检索词",
        },
    ]
    prompt = f"""请为合同审查生成精准的法规/案例检索词，只输出 JSON，不要输出 Markdown。

合同文件名：{fname}
合同类型：{contract_type_name}
用户关注点：{note or "未特别说明"}

合同片段：
{seed}

输出格式：
{{
  "queries": [
    {{"type": "law", "query": "检索词", "reason": "为什么检索"}},
    {{"type": "case", "query": "检索词", "reason": "为什么检索"}}
  ]
}}

要求：
1. 生成 2 个 law 和 2 个 case。
2. query 适合直接送入中文法规/案例数据库检索，避免过长。
3. 聚焦合同争议风险、条款效力、违约责任、解除、付款、管辖等具体问题。
4. 不要编造合同没有体现的专有事实。"""
    try:
        client = get_ai_client()
        response = await client.chat.completions.create(
            model=CONTRACT_REFERENCE_QUERY_MODEL,
            max_tokens=1200,
            temperature=0.2,
            messages=[
                {"role": "system", "content": "你是法律检索关键词生成助手，只输出可解析 JSON。"},
                {"role": "user", "content": prompt},
            ],
        )
        data = _extract_contract_json(response.choices[0].message.content or "{}")
        items = data.get("queries") if isinstance(data, dict) else []
    except Exception:
        return fallback_queries

    queries = []
    seen = set()
    for item in items if isinstance(items, list) else []:
        qtype = str(item.get("type") or "").strip().lower()
        query = re.sub(r"\s+", " ", str(item.get("query") or "")).strip()
        if qtype not in {"law", "case"} or not query:
            continue
        query = query[:300]
        key = (qtype, query)
        if key in seen:
            continue
        seen.add(key)
        queries.append({
            "type": qtype,
            "query": query,
            "reason": str(item.get("reason") or "模型生成的合同审查检索词")[:300],
        })
        if len(queries) >= 4:
            break

    for item in fallback_queries:
        key = (item["type"], item["query"])
        if key not in seen and len([q for q in queries if q["type"] == item["type"]]) < 2:
            queries.append(item)
            seen.add(key)

    return queries or fallback_queries


async def _build_contract_reference_context(
    *,
    fname: str,
    contract_type_name: str,
    note: str,
    description: str,
) -> tuple[str, list[dict]]:
    """按需检索法规/案例依据，返回给模型的上下文和前端展示摘要。"""
    from mcp_client import is_mcp_available, search_law_direct, search_case_direct

    if not is_mcp_available():
        return "（元典检索未启用，未能附加法规/案例依据。）", []

    queries = await _build_contract_reference_queries(
        fname=fname,
        contract_type_name=contract_type_name,
        note=note,
        description=description,
    )

    try:
        results = await asyncio.gather(
            *[
                search_law_direct(item["query"]) if item["type"] == "law" else search_case_direct(item["query"])
                for item in queries
            ]
        )
    except Exception as e:
        return f"（法规/案例检索失败：{e}）", []

    references = []
    for item, result in zip(queries, results):
        references.append({
            "type": item["type"],
            "title": "相关法律法规" if item["type"] == "law" else "相关裁判案例",
            "query": item["query"],
            "reason": item.get("reason", ""),
            "content": str(result or "")[:4000],
        })
    context = "\n\n".join(
        f"【{item['title']}】\n检索词：{item['query']}\n生成理由：{item.get('reason', '')}\n{item['content']}"
        for item in references
        if item["content"].strip()
    )
    return context or "（未检索到可用法规/案例依据。）", references


def _build_contract_review_input(file_info: dict, max_chars: int = 45000) -> tuple[str, dict | None]:
    fname = file_info.get("name", "合同文件")
    fmime = file_info.get("mime_type", "")
    lower_name = fname.lower()
    is_pdf = lower_name.endswith(".pdf") or fmime == "application/pdf"
    is_docx = lower_name.endswith(".docx") or fmime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    if is_pdf or is_docx:
        mineru_error = ""
        doc_label = "PDF" if is_pdf else "Word"
        if os.getenv("MINERU_API_TOKEN", "").strip():
            try:
                parsed = extract_document_with_mineru_package(file_info.get("data", ""), fname)
                markdown = (parsed.get("markdown") or "").strip()
                if markdown:
                    parsed_doc = {
                        "source": "mineru",
                        "document_type": "pdf" if is_pdf else "word",
                        "markdown": markdown[:1000000],
                        "asset_count": parsed.get("asset_count", 0),
                        "has_layout_json": bool(parsed.get("json_files")),
                    }
                    return f"[MinerU解析版{doc_label}: {fname}]\n```\n{markdown[:max_chars]}\n```", parsed_doc
            except Exception as e:
                mineru_error = str(e)
                retry_insecure = os.getenv("MINERU_RETRY_INSECURE_SSL", "true").lower() == "true"
                ssl_error_markers = ("ssl", "certificate", "cert", "tls", "handshake")
                if retry_insecure and any(marker in mineru_error.lower() for marker in ssl_error_markers):
                    try:
                        parsed = extract_document_with_mineru_package(file_info.get("data", ""), fname, verify_ssl=False)
                        markdown = (parsed.get("markdown") or "").strip()
                        if markdown:
                            parsed_doc = {
                                "source": "mineru",
                                "document_type": "pdf" if is_pdf else "word",
                                "markdown": markdown[:1000000],
                                "asset_count": parsed.get("asset_count", 0),
                                "has_layout_json": bool(parsed.get("json_files")),
                            }
                            return f"[MinerU解析版{doc_label}: {fname}]\n```\n{markdown[:max_chars]}\n```", parsed_doc
                    except Exception as retry_error:
                        mineru_error = f"{mineru_error}; SSL 兼容重试失败：{retry_error}"

        fallback = build_file_description(file_info, max_chars=max_chars)
        parsed_doc = {
            "source": "fallback",
            "document_type": "pdf" if is_pdf else "word",
            "markdown": fallback[:120000],
            "error": mineru_error,
            "asset_count": 0,
            "has_layout_json": False,
        }
        return fallback, parsed_doc

    description = build_file_description(file_info, max_chars=max_chars)
    return description, None


def _save_contract_upload_file(user_id: int, file_info: dict) -> str:
    upload_root = Path(os.getenv("CONTRACT_UPLOAD_DIR", Path(__file__).parent / "uploads" / "contract_reviews"))
    upload_root.mkdir(parents=True, exist_ok=True)
    filename = file_info.get("name") or "contract"
    suffix = Path(filename).suffix[:16]
    safe_name = f"{int(time.time())}_{uuid.uuid4().hex}{suffix}"
    user_dir = upload_root / str(user_id)
    user_dir.mkdir(parents=True, exist_ok=True)
    target = user_dir / safe_name
    target.write_bytes(base64.b64decode(file_info.get("data", "")))
    return str(target)


def _split_contract_review_text(text: str, chunk_chars: int, overlap_chars: int = 1200) -> list[str]:
    text = text or ""
    if len(text) <= chunk_chars:
        return [text]
    chunks = []
    start = 0
    while start < len(text):
        end = min(len(text), start + chunk_chars)
        chunks.append(text[start:end])
        if end >= len(text):
            break
        start = max(end - overlap_chars, start + 1)
    return chunks


def _build_contract_review_prompt(
    *,
    fname: str,
    contract_type_name: str,
    contract_type_focus: str,
    note: str,
    include_references: bool,
    reference_context: str,
    contract_text: str,
    issue_limit_instruction: str,
    part_label: str = "",
) -> str:
    scope = f"\n审查范围：{part_label}" if part_label else ""
    return f"""你是资深合同审查律师。请审查下面的合同文本，并只输出 JSON，不要输出 Markdown。

合同文件名：{fname}
合同类型：{contract_type_name}
类型审查重点：{contract_type_focus}
用户关注点：{note or "未特别说明，请做全面风险审查"}{scope}
法规/案例依据开关：{"已开启，请结合下方依据并在 basis 字段中写明可支撑该风险判断的法条、裁判规则或检索依据。" if include_references else "未开启，不需要主动扩展法规/案例检索依据。"}

可参考的法规/案例检索材料：
{reference_context if include_references else "（未启用）"}

合同文本如下，页码可能以“--- 第 N 页 ---”标注：
{contract_text}

请严格输出以下 JSON 结构：
{{
  "summary": "用 2-4 句话概括合同整体风险和最需要关注的问题",
  "overall_risk": "high|medium|low",
  "issues": [
    {{
      "id": "R01",
      "title": "简短问题标题",
      "severity": "high|medium|low",
      "category": "权利义务|付款结算|违约责任|解除终止|争议解决|保密条款|其他",
      "page": 1,
      "quote": "合同中的原文片段，尽量保持原样，便于前端定位",
      "analysis": "为什么这是风险，可能产生什么后果",
      "basis": "可选。开启法规/案例依据时，简要列明相关法条、裁判规则或检索依据；未开启时留空",
      "suggestion": "建议如何修改或补充"
    }}
  ]
}}

要求：
1. {issue_limit_instruction}
2. page 必须尽量根据页码标注填写；无法判断时填 1。
3. quote 必须来自合同原文，不要编造；请截取 15-120 字的连续原文片段，不要改写、拼接或概括。
4. severity 只能是 high、medium、low。
5. 必须输出完整、合法 JSON；如果风险点很多，优先保证 JSON 可解析，不要输出半截内容。
"""


def _combine_contract_review_chunks(chunk_reviews: list[dict], issue_limit: int | None) -> dict:
    risk_rank = {"low": 1, "medium": 2, "high": 3}
    summaries = []
    issues = []
    seen = set()
    overall = "low"
    for review in chunk_reviews:
        summary = str(review.get("summary") or "").strip()
        if summary:
            summaries.append(summary)
        risk = str(review.get("overall_risk") or "medium").lower()
        if risk_rank.get(risk, 2) > risk_rank.get(overall, 1):
            overall = risk
        for item in review.get("issues") or []:
            key = (
                str(item.get("title") or "").strip(),
                str(item.get("quote") or "").strip()[:120],
            )
            if key in seen:
                continue
            seen.add(key)
            issues.append(dict(item))

    if issue_limit and issue_limit > 0:
        order = {"high": 0, "medium": 1, "low": 2}
        issues.sort(key=lambda item: order.get(str(item.get("severity") or "medium").lower(), 1))
        issues = issues[:issue_limit]

    for index, item in enumerate(issues, 1):
        item["id"] = f"R{index:02d}"

    summary = "；".join(summaries[:3]) or "已完成合同分段审查。"
    if len(chunk_reviews) > 1:
        summary = f"已完成 {len(chunk_reviews)} 段合同分段审查。{summary}"
    return {
        "summary": summary[:1200],
        "overall_risk": overall if overall in risk_rank else "medium",
        "issues": issues,
    }


CONTRACT_REVIEW_JOBS: dict[str, dict] = {}


def _public_contract_review_job(job: dict) -> dict:
    payload = {
        "job_id": job["id"],
        "status": job["status"],
        "stage": job.get("stage", ""),
        "message": job.get("message", ""),
        "progress": job.get("progress", 0),
        "filename": job.get("filename", ""),
        "created_at": job.get("created_at"),
        "updated_at": job.get("updated_at"),
    }
    if job.get("result") is not None:
        payload["result"] = job["result"]
    if job.get("error"):
        payload["error"] = job["error"]
    return payload


def _set_contract_review_job(job_id: str, *, status: str | None = None, stage: str | None = None, message: str | None = None, progress: int | None = None, result: dict | None = None, error: str | None = None):
    job = CONTRACT_REVIEW_JOBS.get(job_id)
    if not job:
        return
    if status is not None:
        job["status"] = status
    if stage is not None:
        job["stage"] = stage
    if message is not None:
        job["message"] = message
    if progress is not None:
        job["progress"] = max(0, min(100, int(progress)))
    if result is not None:
        job["result"] = result
    if error is not None:
        job["error"] = error
    job["updated_at"] = time.time()


def _cleanup_contract_review_jobs(max_age_seconds: int = 3600):
    now = time.time()
    stale = [
        job_id for job_id, job in CONTRACT_REVIEW_JOBS.items()
        if now - float(job.get("updated_at") or job.get("created_at") or now) > max_age_seconds
    ]
    for job_id in stale:
        CONTRACT_REVIEW_JOBS.pop(job_id, None)


async def _execute_contract_review(
    *,
    user_id: int,
    session_id: str,
    file_info: dict,
    note: str,
    contract_type: str,
    include_references: bool = False,
    issue_limit: int = 12,
    status_callback=None,
) -> dict:
    """解析合同并生成审查结果；同步接口和后台任务共用。"""
    from db import save_contract_review
    from ai_client import CONTRACT_REVIEW_MODEL, get_ai_client

    async def update(stage: str, message: str, progress: int):
        if status_callback:
            await status_callback(stage, message, progress)

    fname = file_info.get("name", "合同文件")
    file_mime = file_info.get("mime_type", "")
    contract_type_name, contract_type_focus = _get_contract_review_type(contract_type)
    issue_limit = _parse_issue_limit(issue_limit)
    issue_limit_instruction = (
        "不限制风险点数量。请尽可能完整列出所有实质性风险，避免重复、空泛或拆分过细。"
        if issue_limit == 0 else f"最多列出 {issue_limit} 个最重要风险点。"
    )

    await update("parsing", "正在解析合同内容，较大的 PDF/DOCX 可能需要一些时间...", 20)
    max_input_chars = int(os.getenv("CONTRACT_REVIEW_MAX_INPUT_CHARS", "180000"))
    description, parsed_doc = await asyncio.to_thread(_build_contract_review_input, file_info, max_input_chars)
    if parsed_doc is not None:
        parsed_doc["contract_type"] = contract_type if contract_type in CONTRACT_REVIEW_TYPES else "general"
        parsed_doc["contract_type_name"] = contract_type_name
    if "读取失败" in description or "暂不支持" in description:
        raise ValueError(description[:500])

    reference_context = ""
    references = []
    if include_references:
        await update("referencing", "正在检索相关法规和案例依据...", 48)
        reference_context, references = await _build_contract_reference_context(
            fname=fname,
            contract_type_name=contract_type_name,
            note=note,
            description=description,
        )

    await update("reviewing", "合同已解析，正在生成风险审查意见...", 65)
    client = get_ai_client()
    from db import get_user_by_id
    user = await get_user_by_id(user_id)
    chunk_chars = int(os.getenv("CONTRACT_REVIEW_CHUNK_CHARS", "42000"))
    chunks = _split_contract_review_text(description, chunk_chars)
    chunk_reviews = []
    for index, chunk in enumerate(chunks, 1):
        if len(chunks) > 1:
            await update("reviewing", f"合同已解析，正在审查第 {index}/{len(chunks)} 段...", 65 + min(20, int(index / len(chunks) * 20)))
        part_label = f"第 {index}/{len(chunks)} 段。只审查本段文本，风险点 quote 必须来自本段。" if len(chunks) > 1 else ""
        prompt = _build_contract_review_prompt(
            fname=fname,
            contract_type_name=contract_type_name,
            contract_type_focus=contract_type_focus,
            note=note,
            include_references=include_references,
            reference_context=reference_context,
            contract_text=chunk,
            issue_limit_instruction=issue_limit_instruction,
            part_label=part_label,
        )
        response = await client.chat.completions.create(
            model=CONTRACT_REVIEW_MODEL,
            max_tokens=8192,
            temperature=0.2,
            messages=[
                {"role": "system", "content": "你只输出可解析 JSON。不要输出解释、前后缀或 Markdown 代码块。"},
                {"role": "user", "content": prompt},
            ],
        )
        await record_response_usage(
            user=user,
            feature="contract_review",
            model=CONTRACT_REVIEW_MODEL,
            response=response,
        )
        content = response.choices[0].message.content or ""
        review_data = await _extract_or_repair_contract_json(client, content, CONTRACT_REVIEW_MODEL)
        chunk_reviews.append(_normalize_contract_review(review_data, issue_limit=0))

    review = _combine_contract_review_chunks(chunk_reviews, issue_limit=issue_limit)
    if references:
        review["references"] = references

    await update("saving", "审查完成，正在保存记录...", 90)
    file_path = await asyncio.to_thread(_save_contract_upload_file, user_id, file_info)
    record = await save_contract_review(
        user_id=user_id,
        session_id=session_id,
        filename=fname,
        file_mime=file_mime,
        file_path=file_path,
        note=note,
        parsed_doc=parsed_doc,
        review=review,
    )
    return {
        "success": True,
        "filename": fname,
        "review": review,
        "parsed_doc": parsed_doc,
        "record": record,
    }


async def _advance_contract_review_job(job_id: str):
    """在长耗时阶段内持续推进可见进度，避免状态停在排队。"""
    stage_caps = {
        "queued": 10,
        "parsing": 58,
        "referencing": 70,
        "reviewing": 88,
        "saving": 97,
    }
    while True:
        await asyncio.sleep(1)
        job = CONTRACT_REVIEW_JOBS.get(job_id)
        if not job or job.get("status") in ("done", "failed"):
            return

        stage = job.get("stage") or "queued"
        if stage == "queued" and time.time() - float(job.get("created_at") or time.time()) > 1:
            _set_contract_review_job(
                job_id,
                status="running",
                stage="parsing",
                message="正在解析合同内容，较大的 PDF/DOCX 可能需要一些时间...",
                progress=max(int(job.get("progress") or 0), 12),
            )
            continue

        cap = stage_caps.get(stage)
        if cap is None:
            continue
        current = int(job.get("progress") or 0)
        if current < cap:
            step = 2 if stage == "reviewing" else 1
            _set_contract_review_job(job_id, progress=min(cap, current + step))


async def _run_contract_review_job(job_id: str, *, user_id: int, session_id: str, file_info: dict, note: str, contract_type: str, issue_limit: int):
    heartbeat = asyncio.create_task(_advance_contract_review_job(job_id))
    try:
        async def callback(stage: str, message: str, progress: int):
            _set_contract_review_job(job_id, status="running", stage=stage, message=message, progress=progress)

        _set_contract_review_job(job_id, status="running", stage="queued", message="任务已开始，正在准备合同文件...", progress=8)
        result = await _execute_contract_review(
            user_id=user_id,
            session_id=session_id,
            file_info=file_info,
            note=note,
            contract_type=contract_type,
            include_references=bool(CONTRACT_REVIEW_JOBS.get(job_id, {}).get("include_references")),
            issue_limit=issue_limit,
            status_callback=callback,
        )
        _set_contract_review_job(job_id, status="done", stage="done", message="审查完成", progress=100, result=result)
    except Exception as e:
        _set_contract_review_job(job_id, status="failed", stage="failed", message="审查失败", progress=100, error=str(e))
    finally:
        heartbeat.cancel()


@app.post("/api/contract/review")
async def contract_review(request: Request):
    """合同审查工作台：同步解析上传文件并返回结构化风险点。"""
    from db import get_session
    user = await require_current_user(request)
    token = request.cookies.get("session_token")
    session = await get_session(token, user["id"]) if token else None
    if not session:
        raise HTTPException(status_code=401, detail="请先登录")

    try:
        body = await request.json()
        file_info = body.get("file") or {}
        note = (body.get("note") or "").strip()
        contract_type = str(body.get("contract_type") or "general").strip()
        include_references = bool(body.get("include_references", False))
        issue_limit = _parse_issue_limit(body.get("issue_limit", 12))
    except Exception:
        raise HTTPException(status_code=400, detail="无效的请求")

    if not file_info.get("data"):
        raise HTTPException(status_code=400, detail="请上传合同文件")

    try:
        return await _execute_contract_review(
            user_id=user["id"],
            session_id=token,
            file_info=file_info,
            note=note,
            contract_type=contract_type,
            include_references=include_references,
            issue_limit=issue_limit,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"合同审查失败: {str(e)}")


@app.post("/api/contract/review/start")
async def contract_review_start(request: Request):
    """启动合同审查后台任务，前端轮询任务状态。"""
    from db import get_session
    user = await require_current_user(request)
    token = request.cookies.get("session_token")
    session = await get_session(token, user["id"]) if token else None
    if not session:
        raise HTTPException(status_code=401, detail="请先登录")

    try:
        body = await request.json()
        file_info = body.get("file") or {}
        note = (body.get("note") or "").strip()
        contract_type = str(body.get("contract_type") or "general").strip()
        include_references = bool(body.get("include_references", False))
        issue_limit = _parse_issue_limit(body.get("issue_limit", 12))
    except Exception:
        raise HTTPException(status_code=400, detail="无效的请求")

    if not file_info.get("data"):
        raise HTTPException(status_code=400, detail="请上传合同文件")

    _cleanup_contract_review_jobs()
    fname = file_info.get("name", "合同文件")
    job_id = uuid.uuid4().hex
    CONTRACT_REVIEW_JOBS[job_id] = {
        "id": job_id,
        "user_id": user["id"],
        "status": "queued",
        "stage": "queued",
        "message": "任务已提交，等待开始解析...",
        "progress": 3,
        "filename": fname,
        "created_at": time.time(),
        "updated_at": time.time(),
        "result": None,
        "error": "",
        "include_references": include_references,
        "issue_limit": issue_limit,
    }
    asyncio.create_task(_run_contract_review_job(
        job_id,
        user_id=user["id"],
        session_id=token,
        file_info=file_info,
        note=note,
        contract_type=contract_type,
        issue_limit=issue_limit,
    ))
    return {"success": True, "job": _public_contract_review_job(CONTRACT_REVIEW_JOBS[job_id])}


@app.get("/api/contract/review/jobs/{job_id}")
async def contract_review_job_detail(job_id: str, request: Request):
    user = await require_current_user(request)
    job = CONTRACT_REVIEW_JOBS.get(job_id)
    if not job or job.get("user_id") != user["id"]:
        raise HTTPException(status_code=404, detail="审查任务不存在或已过期")
    return {"job": _public_contract_review_job(job)}


@app.post("/api/contract/ask")
async def contract_followup_ask(request: Request):
    """围绕当前合同、解析原文和审查结果进行多轮追问。"""
    from db import get_session
    from ai_client import CONTRACT_REVIEW_MODEL, get_ai_client

    user = await require_current_user(request)
    token = request.cookies.get("session_token")
    session = await get_session(token, user["id"]) if token else None
    if not session:
        raise HTTPException(status_code=401, detail="请先登录")

    try:
        body = await request.json()
        question = str(body.get("question") or "").strip()
        parsed_doc = body.get("parsed_doc") or {}
        review = body.get("review") or {}
        contract_type = str(body.get("contract_type") or parsed_doc.get("contract_type") or "general").strip()
        include_references = bool(body.get("include_references", False))
    except Exception:
        raise HTTPException(status_code=400, detail="无效的请求")

    if not question:
        raise HTTPException(status_code=400, detail="请输入追问内容")

    markdown = str(parsed_doc.get("markdown") or "")[:45000]
    if not markdown:
        raise HTTPException(status_code=400, detail="当前合同没有可用于追问的解析文本")

    contract_type_name, _ = _get_contract_review_type(contract_type)
    reference_context = ""
    references = []
    if include_references:
        reference_context, references = await _build_contract_reference_context(
            fname=str(parsed_doc.get("filename") or "合同文件"),
            contract_type_name=contract_type_name,
            note=question,
            description=markdown[:12000],
        )

    issues = review.get("issues") if isinstance(review, dict) else []
    issues_text = json.dumps(issues[:12] if isinstance(issues, list) else [], ensure_ascii=False, indent=2)
    prompt = f"""你是资深合同审查律师。用户正在围绕同一份合同继续追问。

合同类型：{contract_type_name}
用户问题：{question}

已有审查摘要：
{str(review.get("summary") or "")[:1500] if isinstance(review, dict) else ""}

已有风险点 JSON：
{issues_text[:12000]}

法规/案例依据：
{reference_context if include_references else "（用户未开启依据检索，本次回答不扩展检索。）"}

合同解析原文：
{markdown[:35000]}

请直接回答用户问题。要求：
1. 优先依据当前合同原文和已有审查结果。
2. 如果用户要求修改条款，给出可直接替换的建议文本。
3. 如果开启了依据检索，请简要列明相关法规/案例依据；未开启时不要声称已检索。
4. 不要输出 JSON。"""
    try:
        client = get_ai_client()
        response = await client.chat.completions.create(
            model=CONTRACT_REVIEW_MODEL,
            max_tokens=4096,
            temperature=0.25,
            messages=[
                {"role": "system", "content": "你是一位谨慎、专业的合同审查律师，回答要具体、可操作。"},
                {"role": "user", "content": prompt},
            ],
        )
        await record_response_usage(user=user, feature="contract_followup", model=CONTRACT_REVIEW_MODEL, response=response)
        return {
            "success": True,
            "answer": response.choices[0].message.content or "",
            "references": references,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"追问失败: {str(e)}")


@app.get("/api/contract/reviews")
async def contract_review_list(request: Request):
    from db import list_contract_reviews
    user = await require_current_user(request)
    return {"reviews": await list_contract_reviews(user["id"])}


@app.get("/api/contract/reviews/{review_id}")
async def contract_review_detail(review_id: str, request: Request):
    from db import get_contract_review
    user = await require_current_user(request)
    record = await get_contract_review(review_id, user["id"])
    if not record:
        raise HTTPException(status_code=404, detail="审查记录不存在")
    return {"review": record}


@app.post("/api/contract/reviews/delete")
async def contract_review_delete(request: Request):
    from db import delete_contract_review
    user = await require_current_user(request)
    try:
        body = await request.json()
        review_id = body.get("review_id", "")
    except Exception:
        raise HTTPException(status_code=400, detail="无效的请求")
    if not review_id:
        raise HTTPException(status_code=400, detail="缺少 review_id")
    file_path = await delete_contract_review(review_id, user["id"])
    if file_path is None:
        raise HTTPException(status_code=404, detail="审查记录不存在")
    try:
        upload_root = Path(os.getenv("CONTRACT_UPLOAD_DIR", Path(__file__).parent / "uploads" / "contract_reviews")).resolve()
        target = Path(file_path).resolve()
        if upload_root in target.parents and target.exists():
            target.unlink()
    except Exception:
        pass
    return {"success": True}


@app.get("/api/admin/users")
async def admin_list_users(request: Request):
    await require_admin_user(request)
    from db import list_users
    return {"users": await list_users()}


@app.post("/api/admin/users")
async def admin_create_user(request: Request):
    await require_admin_user(request)
    from db import create_user
    try:
        body = await request.json()
        username = body.get("username", "").strip()
        password = body.get("password", "")
        is_admin = bool(body.get("is_admin", False))
    except Exception:
        raise HTTPException(status_code=400, detail="无效的请求")
    if not username or not password:
        raise HTTPException(status_code=400, detail="用户名和密码不能为空")
    try:
        user = await create_user(username, password, is_admin=is_admin)
        return {"success": True, "user": user}
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/admin/users/update")
async def admin_update_user(request: Request):
    admin = await require_admin_user(request)
    from db import update_user
    try:
        body = await request.json()
        user_id = int(body.get("user_id"))
        password = body.get("password") or None
        is_admin = body.get("is_admin") if "is_admin" in body else None
        is_active = body.get("is_active") if "is_active" in body else None
    except Exception:
        raise HTTPException(status_code=400, detail="无效的请求")

    if user_id == admin["id"] and is_active is False:
        raise HTTPException(status_code=400, detail="不能禁用当前管理员账号")
    if user_id == admin["id"] and is_admin is False:
        raise HTTPException(status_code=400, detail="不能取消当前管理员账号的管理员权限")

    try:
        user = await update_user(
            user_id,
            password=password,
            is_admin=bool(is_admin) if is_admin is not None else None,
            is_active=bool(is_active) if is_active is not None else None,
        )
        return {"success": True, "user": user}
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/api/health")
async def health():
    return {"status": "ok", "version": "1.0.0"}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
