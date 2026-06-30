"""
AI 法律助手 — 文档模板与导出
法律文书模板、python-docx 导出 Word
"""
import hashlib
import os
import re
import tempfile
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 导出目录
EXPORT_DIR = Path(__file__).parent / "exports"

# ── 正则片段 ──────────────────────────────────────────────
_BOLD_ITALIC_RE = re.compile(r"\*\*\*(.+?)\*\*\*")   # ***bold+italic***
_BOLD_RE         = re.compile(r"\*\*(.+?)\*\*")        # **bold**
_ITALIC_RE       = re.compile(r"\*(.+?)\*")            # *italic*
_INLINE_CODE_RE  = re.compile(r"`([^`]+)`")            # `code`
_LINK_RE         = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")  # [text](url)


def _set_run_font(run, name: str = "宋体", size: Optional[Pt] = None,
                  bold: bool = False, italic: bool = False,
                  color: Optional[RGBColor] = None):
    """统一设置 run 字体（含中文字体回退）。"""
    run.font.name = name
    run.element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _append_inline_runs(paragraph, text: str, base_size: Pt = Pt(12)):
    """把一行 Markdown 内联语法拆成多个 run 写入 paragraph。

    策略：从高优先级到低优先级，先把 *** 和 ** 替换成占位符，
    剩下的单 * 就不会误吞多星号内容，最后统一还原。
    """
    # ===== 1. 保护优先处理的内容 =====

    # 1a. 内联代码 `code`
    code_map = {}
    def _code_repl(m):
        ph = f"\x01CODE{len(code_map)}\x01"
        code_map[ph] = m.group(1)
        return ph
    text = _INLINE_CODE_RE.sub(_code_repl, text)

    # 1b. 链接 [text](url)
    link_map = {}
    def _link_repl(m):
        ph = f"\x02LINK{len(link_map)}\x02"
        link_map[ph] = (m.group(1), m.group(2))
        return ph
    text = _LINK_RE.sub(_link_repl, text)

    # ===== 2. 逐层替换星号标记 =====

    # 2a. 先替换 ***bold+italic*** 为占位符
    bi_map = {}
    def _bi_repl(m):
        ph = f"\x03BI{len(bi_map)}\x03"
        bi_map[ph] = m.group(1)
        return ph
    text = _BOLD_ITALIC_RE.sub(_bi_repl, text)

    # 2b. 再替换 **bold** 为占位符（此时不会误吞 ***，因为已被保护）
    b_map = {}
    def _b_repl(m):
        ph = f"\x04B{len(b_map)}\x04"
        b_map[ph] = m.group(1)
        return ph
    text = _BOLD_RE.sub(_b_repl, text)

    # 2c. 最后替换 *italic*（此时不会误伤 ** 或 ***）
    i_map = {}
    def _i_repl(m):
        ph = f"\x05I{len(i_map)}\x05"
        i_map[ph] = m.group(1)
        return ph
    text = _ITALIC_RE.sub(_i_repl, text)

    # ===== 3. 按占位符顺序分段 emit =====
    all_ph = {**code_map, **link_map, **bi_map, **b_map, **i_map}

    remaining = text
    while remaining:
        earliest_ph = None
        earliest_pos = len(remaining)
        for ph in all_ph:
            pos = remaining.find(ph)
            if pos != -1 and pos < earliest_pos:
                earliest_pos = pos
                earliest_ph = ph

        if earliest_ph is None:
            if remaining:
                run = paragraph.add_run(remaining)
                _set_run_font(run, size=base_size)
            break

        if earliest_pos > 0:
            run = paragraph.add_run(remaining[:earliest_pos])
            _set_run_font(run, size=base_size)

        ph = earliest_ph
        value = all_ph[ph]

        if ph in code_map:
            code_run = paragraph.add_run(value)
            _set_run_font(code_run, name="Consolas", size=Pt(base_size.pt * 0.9),
                          color=RGBColor(0xE8, 0x49, 0x6A))
        elif ph in link_map:
            link_text, link_url = value
            link_run = paragraph.add_run(link_text)
            _set_run_font(link_run, size=base_size, bold=True,
                          color=RGBColor(0x25, 0x60, 0xC4))
        elif ph in bi_map:
            run = paragraph.add_run(value)
            _set_run_font(run, size=base_size, bold=True, italic=True)
        elif ph in b_map:
            run = paragraph.add_run(value)
            _set_run_font(run, size=base_size, bold=True)
        elif ph in i_map:
            run = paragraph.add_run(value)
            _set_run_font(run, size=base_size, italic=True)

        remaining = remaining[earliest_pos + len(ph):]


def _set_cell_border(cell, **kwargs):
    """给单元格设置边框。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
        if edge in kwargs:
            element = OxmlElement(f"w:{edge}")
            for attr, val in kwargs[edge].items():
                element.set(qn(f"w:{attr}"), str(val))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def _add_table_borders(table):
    """给整个表格加上统一的细边框。"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    border_style = {"val": "single", "sz": "4", "space": "0", "color": "999999"}
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        for k, v in border_style.items():
            el.set(qn(f"w:{k}"), v)
        borders.append(el)
    tblPr.append(borders)


def _add_shading(cell, color_hex: str):
    """给单元格添加背景色。"""
    tcPr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    tcPr.append(shading)


def _md_to_docx(content: str, title: str) -> Document:
    """
    将 Markdown 内容转换为 Word 文档。
    支持：标题、加粗/斜体、内联代码、代码块、表格、引用、列表、分隔线。
    """
    doc = Document()

    # ── 页面设置：A4 纸张，左侧装订边距 ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.54)

    # ── Normal 样式：宋体 小四（12pt），首行缩进2字符，固定行距28磅 ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    pf = style.paragraph_format
    pf.first_line_indent = Pt(24)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(28)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # ── 文档标题 ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    _set_run_font(title_run, name="黑体", size=Pt(22), bold=True)
    title_para.paragraph_format.first_line_indent = Pt(0)
    title_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    title_para.paragraph_format.space_after = Pt(12)
    doc.add_paragraph()

    lines = content.split("\n")
    i = 0
    in_code_block = False
    code_buf = []
    body_buf = []   # 连续正文行缓存（合并成一个段落）

    def _flush_pending():
        nonlocal in_code_block, code_buf, body_buf
        # 先刷新代码块
        if in_code_block:
            code_text = "\n".join(code_buf)
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(code_text)
            _set_run_font(run, name="Consolas", size=Pt(9),
                          color=RGBColor(0x33, 0x33, 0x33))
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "F2F2F2")
            shd.set(qn("w:val"), "clear")
            pPr.append(shd)
            code_buf = []
            in_code_block = False
        # 再刷新正文缓存（合并为一段）
        if body_buf:
            text = " ".join(body_buf)
            _add_body_paragraph(text, Pt(12))
            body_buf = []

    def _is_table_row(l: str) -> bool:
        return l.startswith("|") and l.rstrip().endswith("|")

    def _is_separator_row(l: str) -> bool:
        return bool(re.match(r"^\|[\s\-:|]+\|$", l))

    def _parse_table_row(l: str) -> list[str]:
        return [c.strip() for c in l.strip("|").split("|")]

    def _add_body_paragraph(text: str, size: Pt = Pt(12), bold: bool = False,
                            font_name: str = "宋体", alignment=None,
                            no_indent: bool = False):
        """添加正文段落（默认首行缩进 + 固定行距）"""
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(28)
        if no_indent:
            p.paragraph_format.first_line_indent = Pt(0)
        else:
            p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if alignment is not None:
            p.alignment = alignment
        _append_inline_runs(p, text, size)
        for r in p.runs:
            r.font.bold = bold
            r.font.name = font_name
            r.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        return p

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        # ── 代码块 ──
        if stripped.startswith("```"):
            if in_code_block:
                code_text = "\n".join(code_buf)
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(code_text)
                _set_run_font(run, name="Consolas", size=Pt(9),
                              color=RGBColor(0x33, 0x33, 0x33))
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "F2F2F2")
                shd.set(qn("w:val"), "clear")
                pPr.append(shd)
                code_buf = []
                in_code_block = False
            else:
                _flush_pending()
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buf.append(lines[i])
            i += 1
            continue

        # ── 空行（刷新正文缓存，输出段落分隔） ──
        if not stripped:
            _flush_pending()
            doc.add_paragraph()
            i += 1
            continue

        # ── 表格 ──
        if _is_table_row(stripped):
            _flush_pending()
            table_lines = []
            while i < len(lines) and _is_table_row(lines[i].rstrip()):
                table_lines.append(lines[i].rstrip().strip())
                i += 1
            if not table_lines:
                continue

            header_row = table_lines[0]
            data_start = 1
            if len(table_lines) > 1 and _is_separator_row(table_lines[1]):
                data_start = 2

            headers = _parse_table_row(header_row)
            if not headers:
                continue
            if all(h == "" for h in headers):
                headers = [f"列{i+1}" for i in range(len(headers))]

            data_rows = [_parse_table_row(r) for r in table_lines[data_start:]]

            num_cols = len(headers)
            for row in data_rows:
                if len(row) < num_cols:
                    row.extend([""] * (num_cols - len(row)))
                elif len(row) > num_cols:
                    row = row[:num_cols]

            num_rows = 1 + len(data_rows)
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            _add_table_borders(table)

            # 表头：深蓝底白字
            for ci, h in enumerate(headers):
                cell = table.cell(0, ci)
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(h)
                _set_run_font(run, name="黑体", size=Pt(10.5), bold=True,
                              color=RGBColor(0xFF, 0xFF, 0xFF))
                _add_shading(cell, "2B5797")

            # 数据行
            for ri, row in enumerate(data_rows):
                for ci, val in enumerate(row):
                    cell = table.cell(ri + 1, ci)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    p.paragraph_format.first_line_indent = Pt(0)
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.space_after = Pt(1)
                    run = p.add_run(val)
                    _set_run_font(run, name="宋体", size=Pt(10.5))
                    if ri % 2 == 1:
                        _add_shading(cell, "F0F4FA")

            doc.add_paragraph()
            continue

        # ── H1 ──
        if stripped.startswith("# ") and not stripped.startswith("## "):
            _flush_pending()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            _append_inline_runs(p, stripped[2:], Pt(18))
            for r in p.runs:
                r.font.bold = True
                r.font.name = "黑体"
                r.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            i += 1
            continue

        # ── H2 ──
        if stripped.startswith("## "):
            _flush_pending()
            _add_body_paragraph(stripped[3:], size=Pt(16), bold=True, font_name="黑体", no_indent=True)
            i += 1
            continue

        # ── H3 ──
        if stripped.startswith("### "):
            _flush_pending()
            _add_body_paragraph(stripped[4:], size=Pt(14), bold=True, font_name="黑体", no_indent=True)
            i += 1
            continue

        # ── H4 ──
        if stripped.startswith("#### "):
            _flush_pending()
            _add_body_paragraph(stripped[5:], size=Pt(13), bold=True, font_name="黑体", no_indent=True)
            i += 1
            continue

        # ── 引用块 ──
        if stripped.startswith("> "):
            _flush_pending()
            quote_lines = []
            while i < len(lines) and lines[i].rstrip().strip().startswith("> "):
                quote_lines.append(lines[i].rstrip().strip()[2:])
                i += 1
            full_quote = " ".join(quote_lines)
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(full_quote)
            _set_run_font(run, name="楷体", size=Pt(11), italic=True,
                          color=RGBColor(0x66, 0x66, 0x66))
            doc.add_paragraph()
            continue

        # ── 分隔线 ──
        if re.match(r"^[-*_]{3,}$", stripped):
            _flush_pending()
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(0)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "CCCCCC")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # ── 有序列表 ──
        if re.match(r"^\d+\.\s", stripped):
            _flush_pending()
            list_items = []
            numbered = True
            while i < len(lines):
                sl = lines[i].rstrip().strip()
                if re.match(r"^\d+\.\s", sl):
                    list_items.append(("num", re.sub(r"^\d+\.\s+", "", sl)))
                    i += 1
                elif sl.startswith("- ") or sl.startswith("* "):
                    numbered = False
                    list_items.append(("bullet", sl[2:]))
                    i += 1
                else:
                    break

            for idx, (kind, item_text) in enumerate(list_items):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.5)
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                p.paragraph_format.line_spacing = Pt(26)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                if numbered and kind == "num":
                    p.style = doc.styles["List Number"]
                    _append_inline_runs(p, item_text, Pt(12))
                else:
                    prefix = f"{idx + 1}. " if kind == "num" else "• "
                    prefix_run = p.add_run(prefix)
                    _set_run_font(prefix_run, size=Pt(12))
                    _append_inline_runs(p, item_text, Pt(12))
            doc.add_paragraph()
            continue

        # ── 无序列表 ──
        if stripped.startswith("- ") or stripped.startswith("* "):
            _flush_pending()
            list_items = []
            while i < len(lines):
                sl = lines[i].rstrip().strip()
                if re.match(r"^\d+\.\s", sl):
                    break
                if sl.startswith("- ") or sl.startswith("* "):
                    list_items.append(sl[2:])
                    i += 1
                elif sl.startswith("  "):
                    if list_items:
                        list_items[-1] += " " + sl.strip()
                    i += 1
                else:
                    break

            for item_text in list_items:
                p = doc.add_paragraph()
                p.style = doc.styles["List Bullet"]
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                p.paragraph_format.line_spacing = Pt(26)
                _append_inline_runs(p, item_text, Pt(12))
            doc.add_paragraph()
            continue

        # ── 普通段落（连续行合并为一段） ──
        body_buf.append(stripped)
        i += 1

    _flush_pending()
    return doc


def export_to_word(content: str, title: str = "法律文书") -> str:
    """
    将内容导出为 Word 文档

    Args:
        content: Markdown 格式的文档内容
        title: 文档标题

    Returns:
        str: 生成的 .docx 文件路径
    """
    EXPORT_DIR.mkdir(exist_ok=True)

    safe_title = re.sub(r"[^\w一-鿿]", "_", title)
    filename = f"{safe_title}_{hashlib.md5(content.encode()).hexdigest()[:8]}.docx"
    filepath = EXPORT_DIR / filename

    doc = _md_to_docx(content, title)
    doc.save(str(filepath))

    return str(filepath)


def cleanup_exports(max_age_hours: int = 24):
    """清理旧的导出文件"""
    import time
    now = time.time()
    if not EXPORT_DIR.exists():
        return
    for f in EXPORT_DIR.glob("*.docx"):
        if now - f.stat().st_mtime > max_age_hours * 3600:
            f.unlink()
